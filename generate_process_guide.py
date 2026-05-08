"""
Sentinel 001 DNV AiP Programme — Compliance Workflow Process Guide
Portrait US Letter, Arial, Navy/White/Light Blue palette
Concise — who/what/when/why/how, community-facing
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY_HEX  = "1F3864"
LBBLUE_HEX = "D5E8F0"
NAVY_RGB  = RGBColor(0x1F, 0x38, 0x64)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_RGB = RGBColor(0x00, 0x00, 0x00)
RED_RGB   = RGBColor(0xC0, 0x00, 0x00)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def set_table_borders(table, color="AAAAAA"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def cell_para(cell, text, bold=False, italic=False, size=10,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

def para(doc, text, size=10, bold=False, italic=False, color=None,
         space_before=4, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT,
         left_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p

def mixed_para(doc, parts, space_before=4, space_after=4,
               align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=None):
    """parts = list of (text, bold, italic, color, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    for text, bold, italic, color, size in parts:
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(size)
        r.font.bold = bold; r.font.italic = italic
        if color: r.font.color.rgb = color
    return p

def section_head(doc, text, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    # Navy bar via shading on paragraph — use a 1-row, 1-col table for the heading bar
    # Actually just style it as bold navy caps
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.bold = True; r.font.color.rgb = NAVY_RGB
    # Add a bottom border to the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def phase_head(doc, number, title, tagline):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    set_table_borders(tbl, color=NAVY_HEX)
    cells = tbl.rows[0].cells
    cells[0].width = Inches(0.55)
    cells[1].width = Inches(5.95)
    shade_cell(cells[0], NAVY_HEX)
    shade_cell(cells[1], LBBLUE_HEX)
    set_cell_margins(cells[0], top=80, bottom=80, left=80, right=80)
    set_cell_margins(cells[1], top=80, bottom=80, left=100, right=100)
    cell_para(cells[0], number, bold=True, color=WHITE_RGB, size=14,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    p = cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r1 = p.add_run(title)
    r1.font.name = "Arial"; r1.font.size = Pt(11); r1.font.bold = True
    r1.font.color.rgb = NAVY_RGB
    r2 = p.add_run(f"  —  {tagline}")
    r2.font.name = "Arial"; r2.font.size = Pt(9); r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)

def bullet(doc, text, bold_prefix=None, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(indent)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = "Arial"; rb.font.size = Pt(10); rb.font.bold = True
        rb.font.color.rgb = NAVY_RGB
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(10)

def add_rule(doc, number, text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    set_table_borders(tbl, color="CCCCCC")
    cells = tbl.rows[0].cells
    cells[0].width = Inches(0.35)
    cells[1].width = Inches(6.15)
    shade_cell(cells[0], NAVY_HEX)
    shade_cell(cells[1], "FFFFFF")
    set_cell_margins(cells[0], top=50, bottom=50, left=60, right=60)
    set_cell_margins(cells[1], top=50, bottom=50, left=100, right=100)
    cell_para(cells[0], number, bold=True, color=WHITE_RGB, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(cells[1], text, size=10)
    doc.add_paragraph().paragraph_format.space_before = Pt(2)

def add_run_field(para, instruction, prefix="", suffix="", size=8):
    if prefix:
        r = para.add_run(prefix); r.font.name = "Arial"; r.font.size = Pt(size)
    run = para.add_run(); run.font.name = "Arial"; run.font.size = Pt(size)
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.text = instruction
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    run._r.extend([fc1, ins, fc2, fc3])
    if suffix:
        r2 = para.add_run(suffix); r2.font.name = "Arial"; r2.font.size = Pt(size)

# ─── Document setup ───────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width    = Inches(8.5)
sec.page_height   = Inches(11)
sec.top_margin    = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin   = Inches(1.0)
sec.right_margin  = Inches(1.0)

doc.styles["Normal"].font.name = "Arial"
doc.styles["Normal"].font.size = Pt(10)

# ─── Header / Footer ──────────────────────────────────────────────────────────
hdr = sec.header
hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]; hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = hp.add_run("Sentinel 001 DNV AiP Programme — Compliance Workflow Process Guide")
r1.font.name = "Arial"; r1.font.size = Pt(8); r1.font.color.rgb = NAVY_RGB; r1.font.bold = True
hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
hp.add_run("\t").font.size = Pt(8)
add_run_field(hp, "PAGE", prefix="Page ", suffix=" of ", size=8)
add_run_field(hp, "NUMPAGES", size=8)

ftr = sec.footer
fp = ftr.paragraphs[0]; fp.clear()
rf = fp.add_run("Bastion Technologies, Inc.  |  CONFIDENTIAL  |  2026-04-26")
rf.font.name = "Arial"; rf.font.size = Pt(7); rf.font.color.rgb = NAVY_RGB

# ─── Title block ──────────────────────────────────────────────────────────────
para(doc, "Sentinel 001 DNV AiP Programme", size=18, bold=True, color=NAVY_RGB,
     space_before=0, space_after=2)
para(doc, "Compliance Workflow Process Guide", size=14, bold=False, color=NAVY_RGB,
     space_before=0, space_after=2)
para(doc, "Bastion Technologies, Inc.  ·  For: Bastion, DEEP, Unique Group, DNV  ·  2026-04-26",
     size=9, italic=True, color=RGBColor(0x55,0x55,0x55), space_before=0, space_after=10)

# Thin navy rule
hr_tbl = doc.add_table(rows=1, cols=1)
hr_tbl.style = "Table Grid"
shade_cell(hr_tbl.rows[0].cells[0], NAVY_HEX)
hr_tbl.rows[0].cells[0].height = Pt(3)
set_cell_margins(hr_tbl.rows[0].cells[0], top=0, bottom=0, left=0, right=0)
doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ─── Purpose ─────────────────────────────────────────────────────────────────
section_head(doc, "Purpose")
para(doc,
     "This document describes how Bastion, DEEP, and DNV work together to advance Sentinel 001 "
     "toward LIVHAB(SAT) classification under DNV-RU-UWT Pt.5 Ch.10 (December 2024). "
     "It covers who does what, when they do it, and why the sequence matters. "
     "It is not a technical specification. It is the operating procedure for the programme team.",
     space_before=4, space_after=6)

# ─── The Three Parties ────────────────────────────────────────────────────────
section_head(doc, "The Three Parties")
para(doc, "", space_before=2, space_after=0)

party_tbl = doc.add_table(rows=4, cols=3)
party_tbl.style = "Table Grid"
set_table_borders(party_tbl, color=NAVY_HEX)
party_tbl.rows[0].cells[0].width = Inches(1.6)
party_tbl.rows[0].cells[1].width = Inches(1.6)
party_tbl.rows[0].cells[2].width = Inches(3.3)

party_data = [
    ("Party", "Who", "Accountable For", True),
    ("Bastion Technologies", "Programme Lead\nLarry Chase / Joe Reeves",
     "All technical documents. Conflict resolution. Submission package. TQ co-ordination. "
     "Driving every action to closure.", False),
    ("DEEP / Unique Group", "Specialist Engineering\nJim Williamson + team",
     "Life support system envelope. CO2, O2, atmospheric monitoring design. "
     "LSS TQ responses. Reporting findings to Bastion immediately — not informally.", False),
    ("DNV", "Classification Society",
     "Document review. Technical Queries. Approval in Principle certificate issuance.", False),
]
for ri, (p1, p2, p3, is_hdr) in enumerate(party_data):
    cells = party_tbl.rows[ri].cells
    bg = NAVY_HEX if is_hdr else ("FFFFFF" if ri % 2 == 0 else LBBLUE_HEX)
    fg = WHITE_RGB if is_hdr else BLACK_RGB
    for ci, txt in enumerate([p1, p2, p3]):
        shade_cell(cells[ci], bg)
        set_cell_margins(cells[ci])
        cell_para(cells[ci], txt, bold=is_hdr, color=fg, size=9 if not is_hdr else 9)

doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ─── The Workflow ─────────────────────────────────────────────────────────────
section_head(doc, "The Workflow")
para(doc, "Five phases. Each one is a prerequisite for the next.",
     size=9, italic=True, space_before=4, space_after=8)

# Phase 1
phase_head(doc, "1", "Document Preparation", "Bastion")
para(doc,
     "Before anything goes to DNV, Bastion generates and internally reviews the full compliance package. "
     "This is not a formatting exercise — it is an engineering review against the rule set.",
     space_before=2, space_after=4)
bullet(doc, "The compliance map (BT3345_9001_009) is built clause by clause against "
       "Pt.5 Ch.10 and cross-checked against the SDD and RTM.")
bullet(doc, "Every conflict between source documents is identified, logged, and given a resolution "
       "owner before DNV sees the package.")
bullet(doc, "Every gap — a requirement not addressed in the design — is explicitly flagged, not buried.")
bullet(doc, "A correction list (BT3345_9001_010) is issued that names every fix, "
       "who owns it, and when it needs to be done.")
para(doc,
     "This phase happens inside Bastion. DNV does not receive documents with unresolved blocking conflicts.",
     bold=True, color=NAVY_RGB, space_before=6, space_after=8)

# Phase 2
phase_head(doc, "2", "Conflict & Gap Resolution", "Bastion + DEEP")
para(doc,
     "Blocking conflicts are engineering disagreements between documents — a spec value in the SDD "
     "that doesn't match the programme baseline. Resolving them requires an engineering decision, "
     "a named owner, and a deadline that feeds the submission schedule.",
     space_before=2, space_after=4)
bullet(doc, "DEEP owns the life support envelope: CO2 setpoints, O2 system design, "
       "atmospheric monitoring thresholds.")
bullet(doc, "When DEEP identifies a non-compliance, Bastion logs it as a programme action immediately. "
       "Findings are not held informally.")
bullet(doc, "Gap items — missing documents, unspecified hardware, incomplete analyses — "
       "each get an owner and a target closure date.")
bullet(doc, "No item is marked resolved until the source document has been updated and reissued.")
para(doc, "", space_before=4, space_after=0)

# Phase 3
phase_head(doc, "3", "DNV Submission", "Bastion → DNV")
para(doc,
     "When blocking items are closed and high-priority items are underway, Bastion submits the "
     "AiP package to DNV. The minimum package comprises:",
     space_before=2, space_after=4)
bullet(doc, "BT3345_9001_009 — Compliance map at current revision")
bullet(doc, "SDD-SENTINEL-001 Rev 1.1 — All blocking conflicts resolved")
bullet(doc, "RTM-SENT-001 — Current approved revision")
bullet(doc, "FMECA-SENTINEL-001 — Complete (10/10 tabs)")
bullet(doc, "PHA-SENTINEL-001 — Issued")
bullet(doc, "BTI-9001-020 — Crew competency submission")
bullet(doc,
       "Submission completeness is tracked in the coordinator's Document Register — "
       "479 documents across 9 subsystems, each showing DNV submission status, "
       "current DNV review status, and open comment count. Verify completeness before submitting.")
para(doc,
     "Nothing goes to DNV that we know is wrong. If we know it is wrong, we fix it first.",
     bold=True, color=NAVY_RGB, space_before=6, space_after=8)

# Phase 4
phase_head(doc, "4", "Technical Query (TQ) Cycle", "DNV → Bastion / DEEP")
para(doc,
     "DNV raises Technical Queries when they need clarification or find issues during review. "
     "Every TQ is a clock that starts running the moment it arrives.",
     space_before=2, space_after=4)
bullet(doc, "Every TQ is logged in the programme coordinator on the day it arrives — "
       "subject, source document, response owner, escalation tier.")
bullet(doc, "DEEP responds to TQs in their domain (life support, atmospheric systems). "
       "Bastion responds to everything else.")
bullet(doc, "Bastion owns every TQ response to DNV, even when the technical content comes from DEEP.")
bullet(doc, "If a TQ cannot be answered within SLA, it is escalated — not left open and unacknowledged. "
       "DNV's time is the critical path.")
bullet(doc,
       "The coordinator's DNV Comments viewer cross-references all Veracity review comments "
       "against the programme action register. Before any TQ response goes to DNV, confirm that "
       "every open 'Action Required' comment on the relevant document has a corresponding "
       "programme action with an owner and due date.")
para(doc, "", space_before=4, space_after=0)

# Phase 5
phase_head(doc, "5", "AiP Certificate", "DNV")
para(doc,
     "DNV issues the Approval in Principle certificate when they are satisfied with the package. "
     "Bastion's job is to make their review as straightforward as possible.",
     space_before=2, space_after=4)
bullet(doc, "The compliance map tells them exactly where we stand — no surprises.")
bullet(doc, "We surface conflicts before they find them.")
bullet(doc, "TQs are responded to faster than the SLA.")
bullet(doc, "The document set is clean, consistent, and internally coherent.")
para(doc,
     "Target: DNV AiP certificate, November 2026.",
     bold=True, color=NAVY_RGB, space_before=6, space_after=10)

# ─── Responsibilities ─────────────────────────────────────────────────────────
section_head(doc, "Responsibilities at a Glance")
para(doc, "", space_before=2, space_after=0)

raci_cols = ["Activity", "Bastion", "DEEP", "DNV"]
raci_data = [
    ("Build compliance map against Pt.5 Ch.10",         "Leads",    "Supports (LSS)", "—"),
    ("Identify conflicts and gaps",                      "Leads",    "Reports to Bastion", "—"),
    ("Resolve conflicts — technical decision",           "Leads",    "Owns LSS domain", "—"),
    ("Issue correction list (BT3345_9001_010)",          "Leads",    "—",              "—"),
    ("Submit AiP package",                               "Leads",    "—",              "Receives"),
    ("Raise Technical Queries",                          "—",        "—",              "Initiates"),
    ("Respond to TQs within SLA",                        "Leads",    "Owns LSS domain", "Receives"),
    ("Log and track all actions to closure",             "Leads",    "Contributes",    "—"),
    ("Issue AiP certificate",                            "—",        "—",              "Issues"),
]

raci_tbl = doc.add_table(rows=len(raci_data)+1, cols=4)
raci_tbl.style = "Table Grid"
set_table_borders(raci_tbl)
raci_widths = [3.0, 1.2, 1.2, 1.1]
hcells = raci_tbl.rows[0].cells
for i, (h, w) in enumerate(zip(raci_cols, raci_widths)):
    shade_cell(hcells[i], NAVY_HEX)
    hcells[i].width = Inches(w)
    set_cell_margins(hcells[i])
    cell_para(hcells[i], h, bold=True, color=WHITE_RGB, size=9,
              align=WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT)

for ri, row_data in enumerate(raci_data):
    cells = raci_tbl.rows[ri+1].cells
    bg = "FFFFFF" if ri % 2 == 0 else LBBLUE_HEX
    for ci, (txt, w) in enumerate(zip(row_data, raci_widths)):
        shade_cell(cells[ci], bg)
        cells[ci].width = Inches(w)
        set_cell_margins(cells[ci])
        is_lead = txt == "Leads"
        cell_para(cells[ci], txt,
                  bold=is_lead, color=NAVY_RGB if is_lead else BLACK_RGB,
                  size=9, align=WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ─── Why We Run It This Way ───────────────────────────────────────────────────
section_head(doc, "Why We Run It This Way")
para(doc,
     "Three things happen in every DNV programme that kill schedules. "
     "This workflow exists to prevent all three.",
     space_before=4, space_after=6)

mixed_para(doc, [
    ("1.  Conflicts are discovered by DNV, not by us.  ", True, False, RED_RGB, 10),
    ("When DNV finds a conflict between your SDD and your RTM, the review stops. "
     "You resubmit. You lose weeks. We find them first, we document them, we resolve them "
     "before the package goes out the door.", False, False, None, 10),
], space_before=4, space_after=4)

mixed_para(doc, [
    ("2.  TQs pile up without owners.  ", True, False, RED_RGB, 10),
    ("A Technical Query with no assigned response owner just sits there. "
     "DNV notices. We assign every TQ a response owner on day one with a date. "
     "That is the only way to keep the review moving.", False, False, None, 10),
], space_before=4, space_after=4)

mixed_para(doc, [
    ("3.  No one knows what's blocking what.  ", True, False, RED_RGB, 10),
    ("When the steel spec conflict is unresolved, the structural calculations can't be finalised. "
     "When the calculations aren't finalised, the plan approval package can't go to DNV. "
     "That chain only becomes visible — and manageable — if someone is tracking it in one place. "
     "The programme coordinator does that — and it now cross-references DNV's own Veracity "
     "comment feed and the full document register across all 9 subsystems, so nothing in "
     "DNV's review queue is invisible to the programme.", False, False, None, 10),
], space_before=4, space_after=6)

para(doc,
     "Running this workflow means Bastion arrives at every DNV interaction prepared. "
     "We know what we have submitted, what is outstanding, what the open questions are, "
     "and who owns each one. That is the difference between a programme that hits its "
     "November 2026 AiP date and one that doesn't.",
     space_before=4, space_after=8)

# ─── Key Rules ────────────────────────────────────────────────────────────────
section_head(doc, "Key Rules")
para(doc, "", space_before=4, space_after=4)

rules = [
    "No document goes to DNV with a known blocking conflict unresolved.",
    "Every TQ gets an owner and a response date on the day it arrives.",
    "DEEP findings are logged as programme actions immediately — not held informally.",
    ("The compliance map (BT3345_9001_009) is the single reference for where we stand "
     "against Pt.5 Ch.10. If it is not in the map, it does not exist as a programme commitment."),
    "Milestone gates are not moved without programme team agreement.",
    ("A conflict is only closed when the source document has been updated, reissued, "
     "and the compliance map has been revised to reflect it."),
    ("Every open 'Action Required' comment in DNV's Veracity register must have a corresponding "
     "programme action in the coordinator with an assigned owner and due date. "
     "The DNV Comments viewer is the cross-reference. Discrepancies are resolved before "
     "the next TQ response goes to DNV."),
]

for i, rule_text in enumerate(rules, 1):
    add_rule(doc, str(i), rule_text)

doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ─── Coordinator Data Reference ──────────────────────────────────────────────
section_head(doc, "Programme Coordinator — Data Reference")
para(doc,
     "The programme coordinator maintains six data registers. All programme-facing data lives here. "
     "DNV Comments and Document Register are read-only views sourced from Veracity and the "
     "programme document management system respectively.",
     space_before=4, space_after=6)

coord_cols = ["Register", "What It Tracks", "Access"]
coord_data = [
    ("Bastion Submissions",      "Document pipeline: PDM approval → DEEP Arena → DNV Veracity. "
                                 "Tracks revision, discipline queue, pipeline status, attachments.",     "Read / Write"),
    ("Joint Actions",            "Programme action register — 94 items with RAG status, owner, "
                                 "due date, and category. Includes seed actions, missing-doc items, "
                                 "and TBC-revision items.",                                              "Read / Write"),
    ("DNV Technical Queries",    "All TQs with 10 business-day SLA tracking, escalation tier (P1–P4), "
                                 "response owner, and routing path.",                                    "Read / Write"),
    ("ISO Audits & Assessments", "DNV independent assessment schedule — readiness score, gate status, "
                                 "required evidence documents, and open TQ blockers.",                   "Read / Write"),
    ("DNV Comments",             "1,271 Veracity review comments across all submitted documents. "
                                 "Filtered by Status (Open / Closed) and Type (Action Required). "
                                 "Read-only cross-reference — do not edit here.",                       "Read-only"),
    ("Document Register",        "479 documents across 9 subsystems. Shows DNV submission status, "
                                 "review status, document completeness, and open comment count.",        "Read-only"),
]

coord_widths = [1.5, 4.2, 0.85]
coord_tbl = doc.add_table(rows=len(coord_data)+1, cols=3)
coord_tbl.style = "Table Grid"
set_table_borders(coord_tbl)
hcells = coord_tbl.rows[0].cells
for i, (h, w) in enumerate(zip(coord_cols, coord_widths)):
    shade_cell(hcells[i], NAVY_HEX)
    hcells[i].width = Inches(w)
    set_cell_margins(hcells[i])
    cell_para(hcells[i], h, bold=True, color=WHITE_RGB, size=9)

for ri, (reg, desc, acc) in enumerate(coord_data):
    cells = coord_tbl.rows[ri+1].cells
    bg = "FFFFFF" if ri % 2 == 0 else LBBLUE_HEX
    for ci, (txt, w) in enumerate(zip([reg, desc, acc], coord_widths)):
        shade_cell(cells[ci], bg)
        cells[ci].width = Inches(w)
        set_cell_margins(cells[ci])
        cell_para(cells[ci], txt, bold=(ci == 0), size=9,
                  color=(NAVY_RGB if ci == 0 else BLACK_RGB))

doc.add_paragraph().paragraph_format.space_before = Pt(6)

para(doc, "PDF Reports — Available for Formal Distribution",
     bold=True, color=NAVY_RGB, size=10, space_before=4, space_after=4)

report_data = [
    ("SLA Forecast",               "Open TQ status against the 10 business-day SLA window. "
                                   "Summary cards (Breached / Critical / Warning / On Track / P1 Open) "
                                   "and full TQ detail table."),
    ("Interface Health Scorecard", "Six KPIs across TQ compliance, actions, submissions, and audit "
                                   "readiness. Discipline breakdown bars and programme snapshot."),
    ("Audit Readiness Forecast",   "Per-audit readiness score, required evidence document status, "
                                   "and open TQ blockers for each upcoming DNV assessment."),
]
rpt_tbl = doc.add_table(rows=len(report_data)+1, cols=2)
rpt_tbl.style = "Table Grid"
set_table_borders(rpt_tbl)
rpt_widths = [2.0, 4.5]
rpt_hdrs = ["Report", "Content"]
hcells = rpt_tbl.rows[0].cells
for i, (h, w) in enumerate(zip(rpt_hdrs, rpt_widths)):
    shade_cell(hcells[i], NAVY_HEX)
    hcells[i].width = Inches(w)
    set_cell_margins(hcells[i])
    cell_para(hcells[i], h, bold=True, color=WHITE_RGB, size=9)

for ri, (name, desc) in enumerate(report_data):
    cells = rpt_tbl.rows[ri+1].cells
    bg = "FFFFFF" if ri % 2 == 0 else LBBLUE_HEX
    for ci, (txt, w) in enumerate(zip([name, desc], rpt_widths)):
        shade_cell(cells[ci], bg)
        cells[ci].width = Inches(w)
        set_cell_margins(cells[ci])
        cell_para(cells[ci], txt, bold=(ci == 0), size=9,
                  color=(NAVY_RGB if ci == 0 else BLACK_RGB))

doc.add_paragraph().paragraph_format.space_before = Pt(6)

# ─── Milestone Schedule ───────────────────────────────────────────────────────
section_head(doc, "Key Milestones")
para(doc, "", space_before=2, space_after=0)

ms_tbl = doc.add_table(rows=6, cols=3)
ms_tbl.style = "Table Grid"
set_table_borders(ms_tbl)
ms_widths = [1.0, 3.0, 2.5]
ms_data = [
    ("Target Date", "Milestone", "Prerequisite", True),
    ("2026-05-15", "Blocking conflicts resolved (Items 001–005 in BT3345_9001_010 closed)",
     "Engineering decisions confirmed by Bastion + DEEP"),
    ("2026-06-12", "SDD-SENTINEL-001 Rev 1.1 issued",
     "All 13 BT3345_9001_010 items closed"),
    ("2026-06-30", "DNV AiP pre-submission technical meeting",
     "All BLOCKING items closed. Key TQs prepared."),
    ("2026-09-30", "AiP package submitted to DNV",
     "All BLOCKING + HIGH items closed. Full document set at approved revisions."),
    ("2026-11-30", "DNV AiP certificate — LIVHAB(SAT)",
     "DNV review complete. All TQs closed."),
]
for ri, row_vals in enumerate(ms_data):
    is_hdr = row_vals[-1] is True if len(row_vals) == 4 else False
    vals = row_vals[:3]
    cells = ms_tbl.rows[ri].cells
    bg = NAVY_HEX if is_hdr else ("FFFFFF" if ri % 2 == 0 else LBBLUE_HEX)
    fg = WHITE_RGB if is_hdr else BLACK_RGB
    for ci, (txt, w) in enumerate(zip(vals, ms_widths)):
        shade_cell(cells[ci], bg)
        cells[ci].width = Inches(w)
        set_cell_margins(cells[ci])
        cell_para(cells[ci], txt, bold=(is_hdr or ci == 0 and not is_hdr),
                  color=(NAVY_RGB if (ci == 0 and not is_hdr) else fg), size=9)

# ─── Save ─────────────────────────────────────────────────────────────────────
out = r"C:\Users\lchas\DNV _System\Sentinel001_AiP_Workflow_Process_Guide.docx"
doc.save(out)
print(f"Saved: {out}")

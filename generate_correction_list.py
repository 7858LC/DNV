"""
BT3345_9001_010 Rev 00 — Sentinel 001 SDD Correction List
python-docx generation script. LibreOffice not installed; DOCX generated natively.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Palette (matches BT3345_9001_009) ────────────────────────────────────────
NAVY_HEX   = "1F3864"
LBBLUE_HEX = "D5E8F0"
WHITE_HEX  = "FFFFFF"
NAVY_RGB   = RGBColor(0x1F, 0x38, 0x64)
WHITE_RGB  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_RGB  = RGBColor(0x00, 0x00, 0x00)

PRIORITY_FMT = {
    "BLOCKING": {"bg": "C00000", "fg": WHITE_RGB},
    "HIGH":     {"bg": "FF8C00", "fg": WHITE_RGB},
    "MEDIUM":   {"bg": "FFC000", "fg": BLACK_RGB},
}
STATUS_FMT = {
    "OPEN": {"bg": "FF0000", "fg": WHITE_RGB},
    "TBD":  {"bg": "808080", "fg": WHITE_RGB},
}

# ─── Helpers (identical pattern to BT3345_9001_009 generator) ─────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def cell_para(cell, text, bold=False, italic=False, size=8,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name   = "Arial"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_run_field(para, instruction, prefix="", suffix="", size=8):
    if prefix:
        r = para.add_run(prefix); r.font.name = "Arial"; r.font.size = Pt(size)
    run = para.add_run(); run.font.name = "Arial"; run.font.size = Pt(size)
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.text = instruction
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    run._r.extend([fc1, ins, fc2, fc3])
    if suffix:
        r2 = para.add_run(suffix); r2.font.name = "Arial"; r2.font.size = Pt(size)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def add_heading_para(doc, text, size=12, bold=True, color=NAVY_RGB,
                     space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text); r.font.name = "Arial"
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    return p

# ─── Correction items ─────────────────────────────────────────────────────────
# Columns: item_no, sdd_section, description, source_doc, conflict_gap_ref,
#          corrective_action, action_owner, priority, status
ITEMS = [
    # ── BLOCKING — Conflicts TBD per programme direction ──────────────────────
    (
        "001",
        "§3.2",
        "Steel material specification mismatch. SDD cites SA-516 Gr.70 (ASME P-No.1, Fy=38 ksi). "
        "Governing programme specification per PDP D000004318 §3.3.1 requires ASTM A576 Gr.50 "
        "(Fy=50 ksi min). Different standard, yield strength, and certification route.",
        "SDD-SENTINEL-001 Rev 1.0 §3.2; PDP D000004318 §3.3.1; RTM-SENT-001 Rev E STR-001",
        "BT3345_9001_009 Row 3.3 CONFLICT; PDP D000004318 §3.3.1",
        "Engineering alignment meeting required. Determine governing specification and update SDD §3.2 "
        "to match. If SA-516 Gr.70 is adopted, update RTM STR-001 accordingly. If A576 Gr.50 is "
        "adopted, confirm ASME equivalence. Flagged TBD pending programme team direction.",
        "Joseph Welker",
        "BLOCKING",
        "TBD",
    ),
    (
        "002",
        "§4.8, §5.3, §6.2",
        "Battery capacity and endurance mismatch. SDD states 40 kWh battery / 80 hr emergency "
        "endurance. Governing programme specification per PDP D000004318 §5.2 requires 60 kWh / "
        "120 hr. SDD provides 67% of required capacity and 67% of required endurance. "
        "Life-safety implication.",
        "SDD-SENTINEL-001 Rev 1.0 §4.8, §5.3, §6.2; PDP D000004318 §5.2; RTM-SENT-001 Rev E BAT-001, BAT-002",
        "BT3345_9001_009 Row 6.3 CONFLICT; PDP D000004318 §5.2",
        "Engineering decision required: upgrade battery bank to 60 kWh or revise programme "
        "specification with documented justification. Update SDD §4.8, §5.3, and §6.2 consistently. "
        "Also update Row 9.2 survival time analysis (see Item 005). Flagged TBD pending programme direction.",
        "Oguzhan Kilic",
        "BLOCKING",
        "TBD",
    ),
    (
        "003",
        "OI Register; §3.5",
        "Escape hatch hinge fatigue — open item not captured in SDD OI register. Hinge fatigue "
        "RPN 175 identified in programme risk register. Fatigue-limited structure (FLS) analysis "
        "not delivered and not referenced in SDD. Cyclic-loaded structural member with no documented "
        "fatigue life assessment.",
        "SDD-SENTINEL-001 Rev 1.0 §3.5, OI Register; Programme risk register",
        "BT3345_9001_009 Row 3.6 GAP; Programme risk register RPN 175",
        "Add escape hatch hinge fatigue analysis as named open item in SDD OI register. "
        "Commission FLS analysis per DNV-RU-UWT Pt.5 Ch.10 Sec.3 [306]. Reference FLS document "
        "in SDD §3.5. Flagged TBD pending FLS delivery.",
        "Joseph Welker / Ken Ngo",
        "BLOCKING",
        "TBD",
    ),
    (
        "004",
        "§2.0 and all affected sections",
        "Operating pressure mismatch. SDD Rev 1.0 states 2.40 bar(a) as operating pressure "
        "throughout. Programme baseline (design basis) is 2.38 bar(a). Discrepancy propagates "
        "to all pressure-dependent calculations including CO2 partial pressure limits, "
        "structural design pressure, and life support sizing.",
        "SDD-SENTINEL-001 Rev 1.0 §2.0, §3.1, §4.4, §5.2, §5.3; Programme design basis document",
        "Programme design basis; see also Item 005 (CO2 pCO2 impact)",
        "Confirm governing design basis operating pressure with programme team. Sweep SDD for all "
        "instances of 2.40 bar(a) and correct to 2.38 bar(a) (or approved value). Re-check all "
        "affected calculations. See Item 013 for systematic sweep. Flagged TBD pending confirmation.",
        "Joseph Welker",
        "BLOCKING",
        "TBD",
    ),
    # ── BLOCKING — Life Safety ────────────────────────────────────────────────
    (
        "005",
        "§4.4",
        "CO2 alarm thresholds stated as percentage by volume, which exceed DNV pCO2 limits at "
        "2.38 bar(a) operating pressure. SDD §4.4: caution 0.5% vol, alarm 1.0% vol. "
        "At 2.38 bar(a): 0.5% vol = 0.0119 bar pCO2 (2.4x DNV normal limit 0.005 bar); "
        "1.0% vol = 0.0238 bar pCO2 (exceeds DNV emergency limit 0.020 bar). "
        "Life-safety non-compliance with DNV-RU-UWT Pt.5 Ch.10 Sec.5 [503].",
        "SDD-SENTINEL-001 Rev 1.0 §4.4; DNV-RU-UWT Pt.5 Ch.10 Sec.5 [503]; RTM-SENT-001 Rev E LSS-003",
        "BT3345_9001_009 Row 5.3 GAP — Correction 1; Jim Williamson (Unique Group) correction actioned",
        "Recalculate CO2 caution and alarm setpoints as pCO2 partial pressure values in bar. "
        "State corrected values in SDD §4.4. Verify against DNV limits: normal <= 0.005 bar pCO2, "
        "emergency <= 0.020 bar pCO2. Update RTM LSS-003 and monitoring system configuration "
        "consistently. Cross-document trace: SDD §4.4, RTM LSS-003, monitoring system config "
        "must all reflect corrected values before DNV submittal.",
        "Jim Williamson (Unique Group)",
        "BLOCKING",
        "OPEN",
    ),
    # ── HIGH ─────────────────────────────────────────────────────────────────
    (
        "006",
        "§1.2",
        "Safety philosophy document (PHA-SENTINEL-001) referenced in SDD §1.2 but not present "
        "in programme document suite as of Rev 00. DNV-RU-UWT Pt.5 Ch.10 Sec.1 [105] requires "
        "formal safety assessment as basis for design. Absence is a compliance gap.",
        "SDD-SENTINEL-001 Rev 1.0 §1.2; DNV-RU-UWT Pt.5 Ch.10 Sec.1 [105]",
        "BT3345_9001_009 Row 1.5 GAP — Correction 2",
        "Initiate PHA-SENTINEL-001. If document exists, locate and add to programme suite. "
        "If not yet issued, remove reference from SDD §1.2 and add as named open item. "
        "PHA must cover all hazard categories per Pt.5 Ch.10 Sec.1 [105] before DNV submittal.",
        "Larry Chase",
        "HIGH",
        "OPEN",
    ),
    (
        "007",
        "§6.4",
        "OI-06 — Battery thermal runaway. Highest hazard classification in SDD. DNV special "
        "approval pathway for lithium battery installations not documented in SDD §6.4. "
        "Battery management system described but thermal runaway mitigation design not specified. "
        "DNV may require Type Approval or special consideration.",
        "SDD-SENTINEL-001 Rev 1.0 §6.4; RTM-SENT-001 Rev E OI-06; DNV-RU-UWT Pt.5 Ch.10 Sec.6 [605]",
        "BT3345_9001_009 Row 6.5 IN PROGRESS; RTM OI-06",
        "Document DNV special approval pathway for lithium battery installation in SDD §6.4. "
        "Include: thermal runaway detection method, suppression or containment approach, "
        "ventilation strategy, and DNV Type Approval certificate reference. "
        "Close RTM OI-06 with completed mitigation design.",
        "Bastion Technologies / Oguzhan Kilic",
        "HIGH",
        "OPEN",
    ),
    (
        "008",
        "§2.2",
        "OI-07 — FMECA full revision required. Structural tab coverage assessed at 3 of 10 "
        "required systems. Dual (redundant) analysis described using SHOULD rather than SHALL, "
        "understating the requirement. FMECA is a prerequisite for DNV plan approval.",
        "SDD-SENTINEL-001 Rev 1.0 §2.2; RTM-SENT-001 Rev E OI-07",
        "BT3345_9001_009 Row 2.4 IN PROGRESS; RTM OI-07",
        "Complete FMECA to cover all 10 required system tabs. Revise SHOULD to SHALL throughout "
        "dual analysis description. FMECA to be submitted to DNV as part of design approval "
        "package. Close RTM OI-07 upon delivery.",
        "Ken Ngo",
        "HIGH",
        "OPEN",
    ),
    (
        "009",
        "§9.6",
        "Rescue mating flange not addressed in SDD Rev 1.0. DNV-RU-UWT Pt.5 Ch.10 Sec.9 [906] "
        "requires standard rescue mating flange for HEV connection. No specification, dimensions "
        "or design referenced. Gap remains pending Flag 4 resolution.",
        "SDD-SENTINEL-001 Rev 1.0 §9.6; RTM-SENT-001 Rev E EMG-006; DNV-RU-UWT Pt.5 Ch.10 Sec.9 [906]",
        "BT3345_9001_009 Row 9.6 GAP; RTM EMG-006",
        "Add rescue mating flange specification to SDD §9.6 per applicable standard. "
        "Confirm compatibility with HEV (see SDD §9.3). Reference standard flange dimensions. "
        "Note: gap to be documented in SDD pending Flag 4 resolution — do not defer without "
        "explicit open item capture.",
        "TBD (pending Flag 4 resolution)",
        "HIGH",
        "OPEN",
    ),
    (
        "010",
        "§3.8",
        "Escape hatch O-ring OEA (Oxygen-Enriched Atmosphere) compatibility — construction hold "
        "item not documented in SDD. O-ring material compatibility in OEA is a safety-critical "
        "material selection issue. Current hold status not captured in SDD OI register.",
        "SDD-SENTINEL-001 Rev 1.0 §3.8; Construction hold log",
        "Construction hold (not yet in compliance map — new finding)",
        "Add O-ring OEA compatibility as named open item in SDD OI register §3.8. "
        "Document material selection criteria, applicable standard (e.g. ASTM G63 / G94), "
        "and current hold status. Resolution required before fabrication of escape hatch.",
        "TBD",
        "HIGH",
        "OPEN",
    ),
    (
        "011",
        "§3.11, §6.1, §6.2, §7.1, §7.3, §10.3",
        "Six COMPLIANT rows in BT3345_9001_009 (Rows 3.12, 6.1, 6.2, 7.1, 7.3, 10.3) are "
        "currently unsupported by cited source documents in SDD. COMPLIANT status was assigned "
        "based on declared design intent; no independent verification by calculation, test or "
        "DNV review has been completed. SDD sections cited for each row do not contain sufficient "
        "documented evidence to sustain COMPLIANT designation.",
        "SDD-SENTINEL-001 Rev 1.0 §3.11, §6.1, §6.2, §7.1, §7.3, §10.3; BT3345_9001_009",
        "BT3345_9001_009 Rows 3.12, 6.1, 6.2, 7.1, 7.3, 10.3 (COMPLIANT — basis required)",
        "For each of the six COMPLIANT rows, add a compliance basis statement to the relevant "
        "SDD section explicitly referencing the calculation, test result or design feature that "
        "demonstrates compliance. Minimum: buoyancy calc ref (§3.11), power load analysis ref "
        "(§6.1), changeover time test ref (§6.2), EMS specification ref (§7.1), "
        "comms system spec ref (§7.3), and this document ref (§10.3).",
        "Larry Chase",
        "HIGH",
        "OPEN",
    ),
    # ── MEDIUM ───────────────────────────────────────────────────────────────
    (
        "012",
        "RTM COVER (not SDD)",
        "RTM-SENT-001 cover page title cell reads 'Rev D | 13 February 2026'. Revision detail "
        "row correctly states Rev E dated 2026-03-20. Title cell is incorrect and constitutes "
        "a document integrity issue that must be corrected before DNV submittal.",
        "RTM-SENT-001 Rev E COVER sheet",
        "RTM cover page discrepancy — flagged in BT3345_9001_009 Phase 1",
        "Correct RTM-SENT-001 cover page title cell to read 'Rev E | 20 March 2026'. "
        "Reissue RTM-SENT-001 Rev E with corrected cover page. Note: this is an RTM action, "
        "not an SDD action, but is included here for programme completeness.",
        "Larry Chase",
        "MEDIUM",
        "OPEN",
    ),
    (
        "013",
        "§2.0 and all affected sections",
        "Operating pressure value 2.40 bar(a) appears throughout SDD Rev 1.0. Pending Conflict 4 "
        "(Item 004) resolution confirming 2.38 bar(a) as governing value, a systematic sweep of "
        "all SDD sections is required to identify and correct every instance.",
        "SDD-SENTINEL-001 Rev 1.0 — all sections; Programme design basis document",
        "Item 004 (CONFLICT 4) — operating pressure sweep action",
        "Following Item 004 resolution, conduct full-text search of SDD for '2.40 bar' and "
        "correct all instances to approved value. Check equations, tables, and figures. "
        "Document sweep results as a revision note in SDD Rev 1.1 change record. "
        "Confirm re-check of all pressure-dependent calculations.",
        "Joseph Welker",
        "MEDIUM",
        "TBD",
    ),
]

print(f"Total items: {len(ITEMS)}")
blocking = sum(1 for i in ITEMS if i[7] == "BLOCKING")
high     = sum(1 for i in ITEMS if i[7] == "HIGH")
medium   = sum(1 for i in ITEMS if i[7] == "MEDIUM")
open_    = sum(1 for i in ITEMS if i[8] == "OPEN")
tbd      = sum(1 for i in ITEMS if i[8] == "TBD")
print(f"BLOCKING: {blocking} | HIGH: {high} | MEDIUM: {medium}")
print(f"OPEN: {open_} | TBD: {tbd}")

# ─── Build document ────────────────────────────────────────────────────────────
doc = Document()

sec0 = doc.sections[0]
sec0.page_width    = Inches(11)
sec0.page_height   = Inches(8.5)
sec0.orientation   = WD_ORIENT.LANDSCAPE
sec0.top_margin    = Inches(0.75)
sec0.bottom_margin = Inches(0.75)
sec0.left_margin   = Inches(0.75)
sec0.right_margin  = Inches(0.75)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(9)

# ─── Page header ──────────────────────────────────────────────────────────────
hdr = sec0.header
hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = hp.add_run(
    "BT3345_9001_010 Rev 00  |  Sentinel 001 SDD Correction List  |  "
    "Input to SDD-SENTINEL-001 Rev 1.1")
r1.font.name = "Arial"; r1.font.size = Pt(8)
r1.font.bold = True; r1.font.color.rgb = NAVY_RGB
hp.paragraph_format.tab_stops.add_tab_stop(Inches(9.5), WD_ALIGN_PARAGRAPH.RIGHT)
hp.add_run("\t").font.size = Pt(8)
add_run_field(hp, "PAGE", prefix="Page ", suffix=" of ", size=8)
add_run_field(hp, "NUMPAGES", size=8)

# ─── Page footer ──────────────────────────────────────────────────────────────
ftr = sec0.footer
fp = ftr.paragraphs[0]
fp.clear()
rf = fp.add_run(
    "PROPRIETARY — Bastion Technologies, Inc.  |  Date: 2026-04-24  |  "
    "Classification: CONFIDENTIAL")
rf.font.name = "Arial"; rf.font.size = Pt(7); rf.font.color.rgb = NAVY_RGB

# ─── Title ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run("SDD Correction List — Sentinel 001 Underwater Habitat")
r.font.name = "Arial"; r.font.size = Pt(16); r.font.bold = True
r.font.color.rgb = NAVY_RGB

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(4)
r2 = p2.add_run(
    "BT3345_9001_010 Rev 00 DRAFT  |  Input to SDD-SENTINEL-001 Rev 1.1  |  2026-04-24")
r2.font.name = "Arial"; r2.font.size = Pt(10); r2.font.color.rgb = NAVY_RGB

# ─── Metadata block ───────────────────────────────────────────────────────────
meta_tbl = doc.add_table(rows=3, cols=4)
meta_tbl.style = "Table Grid"
set_table_borders(meta_tbl)
meta_data = [
    [("Document No.", "BT3345_9001_010"),        ("Revision", "Rev 00 DRAFT")],
    [("Project",      "Sentinel 001 Underwater Habitat — LIVHAB(SAT) AiP Programme"),
                                                  ("Date",     "2026-04-24")],
    [("Owner",        "Larry Chase / Joe Reeves, Bastion Technologies, Inc."),
                                                  ("Status",   "DRAFT — For Review")],
]
for ri, row_data in enumerate(meta_data):
    cells = meta_tbl.rows[ri].cells
    col = 0
    for label, value in row_data:
        shade_cell(cells[col],   NAVY_HEX)
        shade_cell(cells[col+1], LBBLUE_HEX)
        set_cell_margins(cells[col]); set_cell_margins(cells[col+1])
        cell_para(cells[col],   label, bold=True, color=WHITE_RGB, size=8)
        cell_para(cells[col+1], value, size=8)
        col += 2

# ─── Top-of-document note ─────────────────────────────────────────────────────
np = doc.add_paragraph()
np.paragraph_format.space_before = Pt(8)
np.paragraph_format.space_after  = Pt(2)
nr1 = np.add_run("NOTE:  ")
nr1.font.name = "Arial"; nr1.font.size = Pt(8); nr1.font.bold = True
nr1.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
nr2 = np.add_run(
    "All CONFLICT items (Items 001–004) are flagged TBD pending programme team resolution. "
    "Corrected values will not be cited in SDD Rev 1.1 until conflicts are formally resolved "
    "and a governing value is approved by the programme technical authority. "
    "Item 005 (CO2 thresholds) is OPEN and life-safety — correction must not wait on conflict resolution.")
nr2.font.name = "Arial"; nr2.font.size = Pt(8); nr2.font.italic = True

# ─── Priority legend ──────────────────────────────────────────────────────────
add_heading_para(doc, "PRIORITY LEGEND", size=9, space_before=8, space_after=2)
leg_tbl = doc.add_table(rows=1, cols=5)
leg_tbl.style = "Table Grid"
set_table_borders(leg_tbl)
lc = leg_tbl.rows[0].cells
legend_items = [
    ("BLOCKING — Conflict", "C00000", WHITE_RGB),
    ("BLOCKING — Life Safety", "C00000", WHITE_RGB),
    ("HIGH", "FF8C00", WHITE_RGB),
    ("MEDIUM", "FFC000", BLACK_RGB),
    ("OPEN / TBD Status", "808080", WHITE_RGB),
]
for i, (txt, bg, fg) in enumerate(legend_items):
    shade_cell(lc[i], bg)
    set_cell_margins(lc[i], top=30, bottom=30, left=60, right=60)
    cell_para(lc[i], txt, bold=True, color=fg, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_before = Pt(4)

# ─── Main correction table ────────────────────────────────────────────────────
# Col widths (9.5" total):
# Item# | SDD Sec | Description | Source Doc | Gap Ref | Action | Owner | Priority | Status
COL_W = [0.35, 0.65, 2.0, 1.2, 1.1, 2.15, 0.95, 0.65, 0.45]
COL_H = ["Item\nNo.", "SDD\nSection", "Description of Error",
         "Source Document", "Conflict /\nGap Ref",
         "Corrective Action Required", "Action Owner", "Priority", "Status"]

assert abs(sum(COL_W) - 9.5) < 0.01, f"Column widths sum to {sum(COL_W)}"

main_tbl = doc.add_table(rows=1, cols=9)
main_tbl.style = "Table Grid"
set_table_borders(main_tbl)

# Header row
hdr_cells = main_tbl.rows[0].cells
for i, (h, w) in enumerate(zip(COL_H, COL_W)):
    shade_cell(hdr_cells[i], NAVY_HEX)
    hdr_cells[i].width = Inches(w)
    set_cell_margins(hdr_cells[i])
    cell_para(hdr_cells[i], h, bold=True, color=WHITE_RGB, size=8,
              align=WD_ALIGN_PARAGRAPH.CENTER)

# Section dividers
SECTION_BREAKS = {
    "001": ("BLOCKING — CONFLICTS  (TBD pending programme team resolution)", "C00000"),
    "005": ("BLOCKING — LIFE SAFETY",                                        "C00000"),
    "006": ("HIGH PRIORITY",                                                 "FF8C00"),
    "012": ("MEDIUM PRIORITY",                                               "FFC000"),
}

for idx, item in enumerate(ITEMS):
    item_no = item[0]
    priority = item[7]
    status   = item[8]

    # Insert section break row if needed
    if item_no in SECTION_BREAKS:
        label, bg = SECTION_BREAKS[item_no]
        fg = WHITE_RGB if bg != "FFC000" else BLACK_RGB
        br = main_tbl.add_row()
        merged = br.cells[0]
        for j in range(1, 9):
            merged = merged.merge(br.cells[j])
        shade_cell(merged, bg)
        set_cell_margins(merged, top=30, bottom=30, left=80, right=80)
        cell_para(merged, label, bold=True, color=fg, size=9)

    data_row = main_tbl.add_row()
    cells = data_row.cells
    alt = "FFFFFF" if idx % 2 == 0 else "EEF4F8"
    pfmt = PRIORITY_FMT[priority]
    sfmt = STATUS_FMT[status]

    for ci, (val, w) in enumerate(zip(item, COL_W)):
        cells[ci].width = Inches(w)
        set_cell_margins(cells[ci])
        if ci == 7:  # Priority
            shade_cell(cells[ci], pfmt["bg"])
            cell_para(cells[ci], val, bold=True, color=pfmt["fg"], size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        elif ci == 8:  # Status
            shade_cell(cells[ci], sfmt["bg"])
            cell_para(cells[ci], val, bold=True, color=sfmt["fg"], size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            shade_cell(cells[ci], alt)
            is_bold = (ci == 0)
            txt_col = NAVY_RGB if ci == 0 else BLACK_RGB
            cell_para(cells[ci], val, bold=is_bold, color=txt_col, size=8)

# ─── Summary block ────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading_para(doc, "CORRECTION SUMMARY", size=10, space_before=8, space_after=4)

sum_tbl = doc.add_table(rows=3, cols=6)
sum_tbl.style = "Table Grid"
set_table_borders(sum_tbl)

sum_labels = [
    [("Priority", NAVY_HEX, WHITE_RGB), ("BLOCKING (Conflict)", "C00000", WHITE_RGB),
     ("BLOCKING (Life Safety)", "C00000", WHITE_RGB), ("HIGH", "FF8C00", WHITE_RGB),
     ("MEDIUM", "FFC000", BLACK_RGB), ("TOTAL", NAVY_HEX, WHITE_RGB)],
    [("Count", LBBLUE_HEX, NAVY_RGB), (str(blocking-1), LBBLUE_HEX, NAVY_RGB),
     ("1", LBBLUE_HEX, NAVY_RGB), (str(high), LBBLUE_HEX, NAVY_RGB),
     (str(medium), LBBLUE_HEX, NAVY_RGB), (str(len(ITEMS)), LBBLUE_HEX, NAVY_RGB)],
    [("Status", NAVY_HEX, WHITE_RGB), ("OPEN", "FF0000", WHITE_RGB),
     ("", NAVY_HEX, WHITE_RGB), ("TBD", "808080", WHITE_RGB),
     ("", NAVY_HEX, WHITE_RGB), ("", LBBLUE_HEX, NAVY_RGB)],
]

for ri, row_data in enumerate(sum_labels):
    cells = sum_tbl.rows[ri].cells
    for ci, (txt, bg, fg) in enumerate(row_data):
        shade_cell(cells[ci], bg)
        set_cell_margins(cells[ci], top=30, bottom=30, left=60, right=60)
        is_b = (ri == 0 or ci == 0)
        cell_para(cells[ci], txt, bold=is_b, color=fg, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

# Fix row 2 status counts
r2c = sum_tbl.rows[2].cells
shade_cell(r2c[1], "FF0000"); cell_para(r2c[1], str(open_),  bold=True, color=WHITE_RGB, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(r2c[2], LBBLUE_HEX); cell_para(r2c[2], "",        color=NAVY_RGB, size=9)
shade_cell(r2c[3], "808080"); cell_para(r2c[3], str(tbd),    bold=True, color=WHITE_RGB, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_cell(r2c[4], LBBLUE_HEX); cell_para(r2c[4], "",        color=NAVY_RGB, size=9)
shade_cell(r2c[5], LBBLUE_HEX); cell_para(r2c[5], str(len(ITEMS)), bold=True, color=NAVY_RGB, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

# ─── Closing note ─────────────────────────────────────────────────────────────
doc.add_paragraph()
cp = doc.add_paragraph()
cp.paragraph_format.space_before = Pt(6)
cr = cp.add_run(
    "This document is the authoritative input to SDD-SENTINEL-001 Rev 1.1. All 13 items must be "
    "resolved and closed before Rev 1.1 is issued for DNV review. BLOCKING items (001–005) are "
    "prerequisites for DNV plan approval submission. Items 001–004 are flagged TBD; corrected "
    "values will be populated upon programme team direction. Item 005 (CO2 thresholds, life-safety) "
    "is OPEN and must not be deferred. "
    "RTM-SENT-001 cover page discrepancy (Item 012) is an RTM action included for programme "
    "completeness. "
    "Document generated: 2026-04-24 | Next revision upon programme team direction.")
cr.font.name = "Arial"; cr.font.size = Pt(8); cr.font.italic = True

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\lchas\DNV _System\Sentinel001_SDD_Correction_List_Rev00.docx"
doc.save(out_path)
print(f"\nSaved: {out_path}")
print("BT3345_9001_010 Phase complete.")

"""
BT3345_9001_009 Rev 01 — Sentinel 001 DNV Compliance Map
Phase 3: DOCX Generation Script
python-docx required. LibreOffice not used.
All four corrections applied.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colour palette ────────────────────────────────────────────────────────────
NAVY_HEX   = "1F3864"
LBBLUE_HEX = "D5E8F0"
WHITE_HEX  = "FFFFFF"
NAVY_RGB   = RGBColor(0x1F, 0x38, 0x64)
WHITE_RGB  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_RGB  = RGBColor(0x00, 0x00, 0x00)

STATUS_FMT = {
    "COMPLIANT":   {"bg": "00B050", "fg": WHITE_RGB},
    "IN PROGRESS": {"bg": "FFC000", "fg": BLACK_RGB},
    "GAP":         {"bg": "FF0000", "fg": WHITE_RGB},
    "CONFLICT":    {"bg": "C00000", "fg": WHITE_RGB},
    "ASSUMED":     {"bg": "7030A0", "fg": WHITE_RGB},
    "TBD":         {"bg": "808080", "fg": WHITE_RGB},
    "N/A":         {"bg": "D9D9D9", "fg": BLACK_RGB},
}

# ─── Helper functions ──────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def cell_para(cell, text, bold=False, italic=False, size=8,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    run = p.add_run(text)
    run.font.name   = "Arial"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_run_field(paragraph, field_instruction, prefix="", suffix="", size=8):
    if prefix:
        r = paragraph.add_run(prefix)
        r.font.name = "Arial"; r.font.size = Pt(size)
    run = paragraph.add_run()
    run.font.name = "Arial"; run.font.size = Pt(size)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = field_instruction
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
    if suffix:
        r2 = paragraph.add_run(suffix)
        r2.font.name = "Arial"; r2.font.size = Pt(size)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        tblBorders.append(el)
    tblPr.append(tblBorders)

# ─── All 72 compliance rows ────────────────────────────────────────────────────
# Format: (row_num, dnv_ref, req_summary, source_ref, evidence, status, notes)
SECTIONS = [
    {
        "title": "Section 1 — General",
        "rows": [
            ("1.1","Pt.5 Ch.10 Sec.1 [101]",
             "Application — Rules apply to manned underwater habitats for saturation diving operations to 300 m depth.",
             "SDD §1.0; RTM COVER",
             "SDD §1.0 declares Sentinel 001 as a manned saturation diving habitat. Classification target LIVHAB(SAT) stated in RTM COVER. Operating depth 60 m (within 300 m limit).",
             "ASSUMED",
             "Design intent confirmed. Formal DNV scope confirmation required at plan approval stage."),
            ("1.2","Pt.5 Ch.10 Sec.1 [102]",
             "Class notation — LIVHAB(SAT) assigned to habitat meeting all requirements of this chapter.",
             "RTM COVER; SDD §1.0",
             "Classification target LIVHAB(SAT) stated in RTM and SDD. All programme documents reference this notation.",
             "ASSUMED",
             "Full notation assignment pending DNV plan approval. AiP target November 2026."),
            ("1.3","Pt.5 Ch.10 Sec.1 [103]; DNV-OS-E402 (Jul 2024); DNV-RU-OU-0375",
             "Applicable rules — DNV-RU-UWT Pt.5 Ch.10 (December 2024) governs design, construction and operation. DNV-OS-E402 and DNV-RU-OU-0375 are mandatory co-governing standards for LIVHAB(SAT) per Pt.5 Ch.10 Sec.2 [2.1].",
             "SDD §1.0; RTM Sec.1; Programme Context",
             "SDD and RTM reference DNV-RU-UWT Pt.5 Ch.10 Dec 2024 as governing standard. Programme Context confirms: DNV-OS-E402 (July 2024) and DNV-OS-B101 (2024) as supporting standards. Pt.5 Ch.10 Sec.2 [2.1] explicitly requires LIVHAB(SAT) to comply with DNV-OS-E402 and DNV-RU-OU-0375.",
             "ASSUMED",
             "December 2024 edition in force. CO-GOVERNING STANDARDS MANDATORY: DNV-OS-E402 (July 2024) and DNV-RU-OU-0375 must be mapped alongside Pt.5 Ch.10 for LIVHAB(SAT). See Section 12 rows. DNV-OS-B101 (2024) applies as structural material standard. DNV-OS-D201 applies to umbilical systems (RTM SYS-003). Edition control and full co-standard mapping to be confirmed at plan approval."),
            ("1.4","Pt.5 Ch.10 Sec.1 [104]",
             "Definitions — Technical terms as defined in Pt.5 Ch.10 apply throughout design and documentation.",
             "SDD §1.1; RTM Sec.1",
             "No conflicting terminology identified in SDD or RTM. DNV definitions adopted by reference.",
             "ASSUMED",
             "Terminology alignment to be verified during detailed design review."),
            ("1.5","Pt.5 Ch.10 Sec.1 [105]",
             "Safety philosophy — Formal safety assessment (FSA) / PHA required as basis for design. Hazard identification to be systematic and documented.",
             "SDD §1.2; RTM OI Register",
             "No standalone safety philosophy or PHA document identified in programme suite. SDD §1.2 references safety goals but PHA-SENTINEL-001 not located.",
             "GAP",
             "CORRECTION 2 APPLIED. PHA-SENTINEL-001 not present in programme suite as of Rev 00. Action: Larry Chase to initiate formal PHA/FSA document. Verification note: confirm PHA-SENTINEL-001 exists and covers all hazard categories per Pt.5 Ch.10 Sec.1 [105] before DNV submittal."),
            ("1.6","Pt.5 Ch.10 Sec.1 [106]",
             "Risk-based design — Alternative design solutions permissible where equivalence is demonstrated via risk assessment.",
             "SDD §1.2",
             "No alternative design solutions invoked in SDD Rev 1.0. All design solutions follow prescriptive route.",
             "ASSUMED",
             "To be confirmed at detailed design. If alternative solutions are adopted, formal equivalence assessment required."),
            ("1.7","Pt.5 Ch.10 Sec.1 [107]",
             "Deviations — Deviations from rule requirements to be agreed with DNV in writing prior to implementation.",
             "RTM OI Register; SDD §1.2",
             "No deviations formally invoked. Conflicts at Rows 3.3 and 6.3 may require deviation process if not resolved by design update.",
             "ASSUMED",
             "Conflict resolution at Rows 3.3 and 6.3 is prerequisite. Deviation log to be maintained."),
        ]
    },
    {
        "title": "Section 2 — Design Basis & Documentation",
        "rows": [
            ("2.1","Pt.5 Ch.10 Sec.2 [201]",
             "Design basis — Design life, operating depth, working pressure and environmental conditions to be specified and approved.",
             "SDD §2.0; RTM STR-001",
             "SDD §2.0 states: operating depth 60 m, design pressure 7 bar, design life 25 years. Environmental conditions referenced but detailed basis not extracted in Rev 1.0 draft.",
             "ASSUMED",
             "Design basis document or SDD §2.0 to be confirmed against DNV format requirements at plan approval."),
            ("2.2","Pt.5 Ch.10 Sec.2 [202]",
             "Environmental loads — Wave, current, seabed and installation loads to be determined and documented.",
             "SDD §2.1",
             "Environmental load summary referenced in SDD §2.1. Specific load calculations not included in Rev 1.0 draft.",
             "ASSUMED",
             "Full environmental load calculations required as separate submission, to be included in calculations package."),
            ("2.3","Pt.5 Ch.10 Sec.2 [203]",
             "Operating pressure — Internal, external hydrostatic and collapse pressures to be calculated for all operational modes.",
             "SDD §3.1; RTM STR-001",
             "SDD §3.1 references pressure vessel design at 60 m depth. Collapse pressure calculations referenced but not extracted from draft.",
             "ASSUMED",
             "Pressure calculations to be submitted as part of structural calculations package."),
            ("2.4","Pt.5 Ch.10 Sec.2 [204]",
             "FMECA — Failure mode and effects criticality analysis covering all safety-critical systems required before design freeze.",
             "RTM OI-07; SDD §2.2",
             "RTM OI-07 records FMECA structural tab as open item. SDD §2.2 references FMECA requirement. No completed FMECA in programme suite.",
             "IN PROGRESS",
             "OI-07 open. FMECA structural tab in progress. Action owner per RTM. Completion required before DNV plan approval."),
            ("2.5","Pt.5 Ch.10 Sec.2 [205]",
             "HAZID — Hazard identification study to be conducted and recorded with mitigation measures identified.",
             "SDD §1.2; RTM OI Register",
             "HAZID referenced in SDD §1.2 as planned activity. No completed HAZID report in programme suite at Rev 00.",
             "ASSUMED",
             "HAZID to be completed and submitted. Recommend combining with PHA per Row 1.5."),
            ("2.6","Pt.5 Ch.10 Sec.2 [206]",
             "Operational philosophy — Intended operating profile, crew complement, saturation depth and duration to be documented.",
             "SDD §2.3",
             "SDD §2.3 states: crew complement 2 (up to 4 in emergency), saturation depth 60 m, maximum saturation duration 28 days.",
             "ASSUMED",
             "Operational philosophy to be formally documented per DNV submission format."),
        ]
    },
    {
        "title": "Section 3 — Structural Design",
        "rows": [
            ("3.1","Pt.5 Ch.10 Sec.3 [301]",
             "Structural analysis — FEA or classical pressure vessel analysis required. Safety factors per applicable standard.",
             "SDD §3.0",
             "SDD §3.0 states structural analysis per ASME BPVC Sec.VIII Div.1. FEA referenced as planned activity.",
             "ASSUMED",
             "FEA calculations to be included in structural calculations package for DNV review."),
            ("3.2","Pt.5 Ch.10 Sec.3 [302]",
             "Pressure vessel standard — Habitat pressure shell to be designed to ASME BPVC or equivalent DNV-approved standard.",
             "SDD §3.0; RTM STR-001",
             "SDD §3.2 references ASME BPVC Sec.VIII Div.1 as design standard. RTM STR-001 references ASME equivalence.",
             "ASSUMED",
             "Standard applicability to be confirmed with DNV. Material conflict at Row 3.3 affects this assessment."),
            ("3.3","Pt.5 Ch.10 Sec.3 [303]",
             "Structural material — Steel grade to meet notch toughness and yield strength for operating temperature and pressure.",
             "SDD §3.2; RTM STR-001",
             "SDD §3.2: SA-516 Gr.70 (ASME P-No.1, Fy=38 ksi, UTS=70–90 ksi). RTM STR-001: A576 Gr.50 (Fy=50 ksi min, UTS=70 ksi min). CONFLICT: different standards and yield strengths.",
             "CONFLICT",
             "BLOCKING CONFLICT. SDD: SA-516 Gr.70. RTM: A576 Gr.50. Action: Engineering alignment meeting required. SDD Rev 1.1 or RTM Rev F must resolve before DNV submittal. Owner: Larry Chase / lead structural engineer."),
            ("3.4","Pt.5 Ch.10 Sec.3 [304]",
             "Weld procedure qualification — All structural welds to be qualified per applicable standard. NDE requirements to be specified.",
             "SDD §3.3",
             "SDD §3.3 references weld procedure qualification per AWS D1.1 / ASME IX. NDE plan referenced as planned document.",
             "ASSUMED",
             "WPS/PQR and NDE plan to be submitted. Material resolution at Row 3.3 required first."),
            ("3.5","Pt.5 Ch.10 Sec.3 [305]",
             "Pressure testing — Hydrostatic test at 1.5x design pressure required after fabrication. Test plan to be submitted to DNV.",
             "SDD §3.4; RTM STR-003",
             "SDD §3.4 references hydrostatic test at 1.5x design pressure (approx. 10.5 bar for 7 bar design). RTM STR-003 records requirement.",
             "ASSUMED",
             "Test plan to be submitted at fabrication stage. Test to be witnessed by DNV surveyor."),
            ("3.6","Pt.5 Ch.10 Sec.3 [306]",
             "Fatigue analysis — Fatigue assessment required for cyclic-loaded members including hinges, trunnions and pressure connections.",
             "SDD §3.5",
             "SDD §3.5 references fatigue assessment as planned. Hinge fatigue analysis not present in Rev 1.0. No fatigue calculations in programme suite.",
             "GAP",
             "Hinge fatigue analysis absent from SDD Rev 1.0 and programme suite. Gap identified in Phase 2. Action: Commission hinge fatigue analysis. Required before DNV plan approval."),
            ("3.7","Pt.5 Ch.10 Sec.3 [307]",
             "Corrosion protection — Sacrificial anodes or ICCP system required. Coating specification to be submitted.",
             "SDD §3.6",
             "SDD §3.6 references cathodic protection system and approved coating system. Details referenced but not extracted from draft.",
             "ASSUMED",
             "Corrosion protection specification to be included in material and coating documentation package."),
            ("3.8","Pt.5 Ch.10 Sec.3 [308]",
             "Penetrations & hatches — All pressure hull penetrations to maintain structural integrity. Access hatch design loads to be specified.",
             "SDD §3.7; RTM STR-002",
             "SDD §3.7 references penetration design per ASME. RTM STR-002 records hatch design requirement.",
             "ASSUMED",
             "Penetration and hatch calculations to be included in structural package."),
            ("3.9","Pt.5 Ch.10 Sec.3 [309]",
             "Viewports — Viewport design to ASME PVHO-1 or equivalent. Design pressure, impact resistance and light transmittance to be specified.",
             "SDD §3.8",
             "SDD §3.8 references viewport design to ASME PVHO-1. Material (acrylic) and dimensions referenced.",
             "ASSUMED",
             "Viewport calculations and material test certificates required at plan approval."),
            ("3.10","Pt.5 Ch.10 Sec.3 [310]",
             "Lifting & handling — Lift points to be designed for static lift load x 2.5 safety factor. Rigging plan required.",
             "SDD §3.9",
             "SDD §3.9 references lifting pad eyes designed to 2.5x static load. Rigging plan referenced as planned document.",
             "ASSUMED",
             "Lift point calculations and rigging plan to be submitted for DNV review."),
            ("3.11","Pt.5 Ch.10 Sec.3 [311]",
             "Anchor / mooring — Seabed anchor or mooring to be designed for all environmental load combinations.",
             "SDD §3.10",
             "SDD §3.10 references anchor piles / seabed frame design. Detailed anchor calculations not extracted from Rev 1.0.",
             "ASSUMED",
             "Anchor/mooring calculations required as part of structural package."),
            ("3.12","Pt.5 Ch.10 Sec.3 [312]",
             "Buoyancy — Positive buoyancy margin to be demonstrated for all operational conditions including maximum payload.",
             "SDD §3.11; RTM STR-004",
             "SDD §3.11 documents buoyancy calculation showing 5% positive margin at maximum payload. RTM STR-004 status: Verified.",
             "COMPLIANT",
             "Buoyancy calculation confirmed. 5% positive margin demonstrated at maximum payload. RTM STR-004: Verified."),
        ]
    },
    {
        "title": "Section 4 — Compartmentation & Systems",
        "rows": [
            ("4.1","Pt.5 Ch.10 Sec.4 [401]",
             "Watertight integrity — All openings in pressure boundary to maintain watertight integrity during normal and emergency operations.",
             "SDD §4.1; RTM SYS-001",
             "SDD §4.1 describes pressure hull integrity provisions. Penetrations sealed to ASME standard. RTM SYS-001 records requirement.",
             "ASSUMED",
             "Watertight integrity to be verified by pressure test (Row 3.5). Opening schedules to be included in documentation."),
            ("4.2","Pt.5 Ch.10 Sec.4 [402]",
             "Bilge system — Bilge pump or drainage system required for any enclosed wet spaces.",
             "SDD §4.2",
             "Moon-pool design of Sentinel 001 creates ambiguity regarding bilge system applicability. SDD §4.2 does not explicitly address bilge for open-bottom configuration.",
             "TBD",
             "Bilge system applicability unclear for moon-pool habitat. Action: Seek DNV technical clarification on Pt.5 Ch.10 Sec.4 [402] applicability to open-bottom habitat. Include in AiP pre-submission agenda."),
            ("4.3","Pt.5 Ch.10 Sec.4 [403]",
             "Flooding protection — Sensors and alarms to detect water ingress into any enclosed space.",
             "SDD §4.3; RTM SYS-002",
             "SDD §4.3 references water ingress sensors in all enclosed compartments with alarm to surface and habitat. RTM SYS-002 records requirement.",
             "ASSUMED",
             "Sensor locations and alarm setpoints to be specified in system design documents."),
            ("4.4","Pt.5 Ch.10 Sec.4 [404]",
             "Fire detection & suppression — Automatic fire detection in all manned spaces. Suppression appropriate to risk.",
             "SDD §4.4; RTM SYS-003",
             "SDD §4.4 references smoke detectors and fire suppression (halon-free agent). RTM SYS-003 records requirement.",
             "ASSUMED",
             "Fire system design to be submitted. Suppression agent selection critical in pressurised atmosphere — verify agent suitability."),
            ("4.5","Pt.5 Ch.10 Sec.4 [405]",
             "Pressure relief — Pressure relief valve(s) to protect against over-pressurisation of any pressure vessel or system.",
             "SDD §4.5; RTM SYS-004",
             "SDD §4.5 references PRVs on all pressure systems with specified set pressures. RTM SYS-004 records requirement.",
             "ASSUMED",
             "PRV sizing and set pressure calculations required. Test certificates to accompany fabrication."),
            ("4.6","Pt.5 Ch.10 Sec.4 [406]",
             "Umbilical penetrations — Umbilical entry point to be sealed to prevent ingress and maintain atmospheric integrity.",
             "SDD §4.6; RTM SYS-005",
             "SDD §4.6 references umbilical stuffing box design. RTM SYS-005 records requirement.",
             "ASSUMED",
             "Umbilical penetration design to be included in structural package."),
        ]
    },
    {
        "title": "Section 5 — Life Support Systems",
        "rows": [
            ("5.1","Pt.5 Ch.10 Sec.5 [501]",
             "O2 supply — Primary and secondary (emergency) oxygen supply systems required. Capacity for minimum 24 hours at maximum crew.",
             "SDD §5.1; RTM LSS-001, LSS-006",
             "SDD §5.1 references primary O2 supply via surface umbilical and secondary O2 cylinder bank. RTM LSS-006: O2 solenoid normally closed (fail-safe, safety-critical designation confirmed).",
             "ASSUMED",
             "O2 system design to be submitted. RTM LSS-006 safety-critical designation (normally closed solenoid) to be maintained in all design revisions."),
            ("5.2","Pt.5 Ch.10 Sec.5 [502]",
             "CO2 scrubbing — CO2 removal system to maintain pCO2 <= 0.005 bar normal, <= 0.020 bar emergency.",
             "SDD §5.2; RTM LSS-002",
             "SDD §5.2 references CO2 scrubber (Sofnolime). Scrubber capacity and canister duration referenced. pCO2 limits apply per Sec.5 [503].",
             "ASSUMED",
             "Scrubber capacity calculation to be submitted. Cross-reference Row 5.3 for alarm threshold gap."),
            ("5.3","Pt.5 Ch.10 Sec.5 [503]",
             "CO2 alarm thresholds — CO2 monitoring to alarm at <= 0.005 bar pCO2 (normal) and <= 0.020 bar pCO2 (emergency).",
             "SDD §4.4; RTM LSS-003; DNV Pt.5 Ch.10 Sec.5 [503]",
             "SDD §4.4 sets CO2 alarms at 0.5% vol (caution) and 1.0% vol (alarm). At operating pressure 2.38-2.40 bar(a): 0.5% vol = 0.0119-0.012 bar pCO2 (2.4x DNV normal limit 0.005 bar); 1.0% vol = 0.0238-0.024 bar pCO2 (exceeds DNV emergency limit 0.020 bar).",
             "GAP",
             "CORRECTION 1 APPLIED. pCO2 arithmetic confirms SDD §4.4 % volume thresholds non-compliant with DNV Pt.5 Ch.10 Sec.5 [503] at operating pressure. Jim Williamson (Unique Group) notified and has corrected setpoint value. Cross-document trace flag: SDD §4.4, RTM LSS-003, and monitoring system configuration must all be updated to pCO2-based limits. Action owner: Larry Chase / Jim Williamson."),
            ("5.4","Pt.5 Ch.10 Sec.5 [504]",
             "Atmospheric monitoring — Continuous monitoring of O2, CO2, CO and hydrocarbons with audible and visual alarms in all manned spaces.",
             "SDD §5.3; RTM LSS-004",
             "SDD §5.3 references continuous atmospheric monitoring for O2, CO2, CO and hydrocarbons. Alarm system covers all manned spaces.",
             "ASSUMED",
             "Atmospheric monitoring system specification to be submitted. CO2 threshold gap addressed at Row 5.3."),
            ("5.5","Pt.5 Ch.10 Sec.5 [505]",
             "Temperature & humidity — Habitat temperature 18-26 degrees C. Relative humidity <= 85%.",
             "SDD §5.4; RTM LSS-005",
             "SDD §5.4 references HVAC maintaining temperature 20 +/- 2 degrees C and humidity <80% RH. Both within DNV limits.",
             "ASSUMED",
             "HVAC system design and capacity calculations to be submitted."),
            ("5.6","Pt.5 Ch.10 Sec.5 [506]",
             "Potable water — Minimum 2 litres per person per day for saturation duration. Emergency reserve to be identified.",
             "SDD §5.5; RTM LSS-007",
             "SDD §5.5 references surface umbilical water supply and 5-day emergency reserve storage on habitat. RTM LSS-007 records requirement.",
             "ASSUMED",
             "Water capacity calculation to be verified against crew complement (4) and maximum saturation duration (28 days)."),
            ("5.7","Pt.5 Ch.10 Sec.5 [507]",
             "Sanitary systems — Waste management provisions required. No overboard discharge of raw sewage.",
             "SDD §5.6; RTM LSS-008",
             "SDD §5.6 references vacuum toilet system with holding tank. No overboard discharge. RTM LSS-008 records requirement.",
             "ASSUMED",
             "Holding tank capacity to be calculated for maximum saturation duration (28 days, crew 4)."),
            ("5.8","Pt.5 Ch.10 Sec.5 [508]",
             "Medical oxygen — Dedicated medical O2 supply separate from atmospheric system. Sufficient for emergency recompression.",
             "SDD §5.7; RTM LSS-009",
             "SDD §5.7 references medical O2 cylinder bank, separate from primary atmospheric supply.",
             "ASSUMED",
             "Medical O2 capacity calculation and cylinder certification to be submitted."),
        ]
    },
    {
        "title": "Section 6 — Power & Electrical Systems",
        "rows": [
            ("6.1","Pt.5 Ch.10 Sec.6 [601]",
             "Main power supply — Primary power from surface via umbilical. Capacity sufficient for all normal operating loads.",
             "SDD §6.1; RTM ELE-001",
             "SDD §6.1 documents surface umbilical power at 440V 3-phase. Load analysis shows 22 kW normal load. RTM ELE-001 status: Verified.",
             "COMPLIANT",
             "Main power supply confirmed compliant. Surface umbilical power verified for normal operations."),
            ("6.2","Pt.5 Ch.10 Sec.6 [602]",
             "Emergency power — Automatic changeover to emergency power within 45 seconds of main supply failure.",
             "SDD §6.2; RTM ELE-002",
             "SDD §6.2 documents automatic changeover to battery UPS in <10 seconds. RTM ELE-002 status: Verified.",
             "COMPLIANT",
             "Emergency power changeover confirmed compliant. <10 s changeover exceeds 45 s rule requirement."),
            ("6.3","Pt.5 Ch.10 Sec.6 [603]",
             "Battery capacity & endurance — Emergency battery to sustain life-critical loads for minimum 72 hours.",
             "SDD §4.8, §5.3; RTM BAT-001, BAT-002",
             "SDD §4.8 and §5.3: 40 kWh / 80 hr endurance. RTM BAT-001: 60 kWh required. RTM BAT-002: 120 hr endurance required. CONFLICT: SDD provides 67% of required capacity and 67% of required endurance.",
             "CONFLICT",
             "BLOCKING CONFLICT. SDD: 40 kWh / 80 hr. RTM: 60 kWh / 120 hr. Deficiency: 20 kWh and 40 hr. Action: Upgrade battery bank to 60 kWh or revise RTM with justification. SDD Rev 1.1 or RTM Rev F must resolve. Owner: Larry Chase / electrical engineer."),
            ("6.4","Pt.5 Ch.10 Sec.6 [604]",
             "Electrical distribution — Segregated distribution for safety-critical and non-critical loads. Protection co-ordination study required.",
             "SDD §6.3; RTM ELE-003",
             "SDD §6.3 references segregated distribution boards for safety-critical loads. Protection study referenced as planned.",
             "ASSUMED",
             "Electrical distribution schematic and protection co-ordination study to be submitted."),
            ("6.5","Pt.5 Ch.10 Sec.6 [605]",
             "Battery thermal runaway — Detection and mitigation measures required for lithium battery installations.",
             "RTM OI-06; SDD §6.4",
             "RTM OI-06 records battery thermal runaway protection as open item. SDD §6.4 references BMS but detailed thermal runaway mitigation not specified in Rev 1.0.",
             "IN PROGRESS",
             "OI-06 open. Battery thermal runaway analysis and mitigation design in progress. Action owner per RTM. Required before DNV plan approval."),
            ("6.6","Pt.5 Ch.10 Sec.6 [606]",
             "Cable routing & protection — Cable selection, routing and fire protection in manned spaces per applicable standard.",
             "SDD §6.5; RTM ELE-004",
             "SDD §6.5 references LSOH cabling in manned spaces. Routing and protection referenced.",
             "ASSUMED",
             "Cable schedule and routing drawings to be submitted. LSOH compliance to be certified."),
            ("6.7","Pt.5 Ch.10 Sec.6 [607]",
             "Lighting — Normal and emergency lighting (battery-backed) in all manned spaces. Minimum illumination levels to be met.",
             "SDD §6.6; RTM ELE-005",
             "SDD §6.6 references normal LED and emergency lighting with battery backup. Illumination levels referenced.",
             "ASSUMED",
             "Lighting layout and illumination calculations to be submitted."),
            ("6.8","Pt.5 Ch.10 Sec.6 [608]",
             "Earthing & bonding — All metallic components to be bonded. Earth fault detection for IT systems.",
             "SDD §6.7; RTM ELE-006",
             "SDD §6.7 references earthing and bonding scheme with earth fault detection.",
             "ASSUMED",
             "Earthing and bonding schedule to be submitted."),
        ]
    },
    {
        "title": "Section 7 — Control & Monitoring Systems",
        "rows": [
            ("7.1","Pt.5 Ch.10 Sec.7 [701]",
             "Environmental monitoring system — Continuous real-time monitoring of all life support parameters with data logging.",
             "SDD §7.1; RTM MON-001",
             "SDD §7.1 documents integrated EMS with real-time monitoring of O2, CO2, CO, temperature, humidity, pressure and depth. Data logging to surface and local storage. RTM MON-001: Verified.",
             "COMPLIANT",
             "Environmental monitoring system confirmed compliant with continuous monitoring and data logging."),
            ("7.2","Pt.5 Ch.10 Sec.7 [702]",
             "Alarm management — Alarm hierarchy (caution/warning/alarm/emergency) with priority routing to habitat and surface.",
             "SDD §7.2; RTM MON-002",
             "SDD §7.2 references four-tier alarm hierarchy with priority routing. Alarm rationalisation referenced as planned.",
             "ASSUMED",
             "Alarm management philosophy and rationalisation study to be submitted. CO2 threshold gap addressed at Row 5.3."),
            ("7.3","Pt.5 Ch.10 Sec.7 [703]",
             "Communication systems — Two independent voice communication systems between habitat and surface. Diver communication required.",
             "SDD §7.3; RTM COM-001",
             "SDD §7.3 documents primary hardwired and secondary acoustic communication systems. Diver through-water communication included. RTM COM-001: Verified.",
             "COMPLIANT",
             "Communication systems confirmed compliant. Two independent systems with diver comms verified."),
            ("7.4","Pt.5 Ch.10 Sec.7 [704]",
             "Data recording — All critical parameters to be recorded with minimum 30-day retention. Time-stamped records.",
             "SDD §7.4; RTM MON-003",
             "SDD §7.4 references data recording with 90-day local storage and surface backup. RTM MON-003 records requirement.",
             "ASSUMED",
             "Data recording system specification to be submitted. 90-day retention exceeds 30-day minimum."),
            ("7.5","Pt.5 Ch.10 Sec.7 [705]",
             "Remote monitoring — Surface support vessel or control room to have real-time visibility of all critical habitat parameters.",
             "SDD §7.5; RTM MON-004",
             "SDD §7.5 references surface monitoring console with real-time parameter display and alarm reception.",
             "ASSUMED",
             "Surface monitoring system to be specified and tested during commissioning."),
            ("7.6","Pt.5 Ch.10 Sec.7 [706]",
             "Control system redundancy — Safety-critical control functions to have redundant execution paths. Single failure not to cause loss of life support.",
             "SDD §7.6; RTM MON-005",
             "SDD §7.6 references redundant control architecture for life support and emergency systems. Details to be confirmed in detailed design.",
             "ASSUMED",
             "Control system FMECA to be included in overall FMECA (OI-07). Redundancy to be demonstrated in system design."),
        ]
    },
    {
        "title": "Section 8 — Diving Equipment & Access",
        "rows": [
            ("8.1","Pt.5 Ch.10 Sec.8 [801]",
             "Diver entry/exit — Moon-pool or lock-out arrangements to allow diver entry and exit under pressure.",
             "SDD §8.1; RTM DIV-001",
             "SDD §8.1 describes moon-pool entry configuration. Pressurised water column at operating depth allows diver entry/exit without hull penetration.",
             "ASSUMED",
             "Moon-pool configuration to be reviewed against DNV requirements. Interface with bilge TBD at Row 4.2."),
            ("8.2","Pt.5 Ch.10 Sec.8 [802]",
             "Diver equipment stowage — Stowage for diving equipment including bailout gas, helmets and suits for full crew complement.",
             "SDD §8.2; RTM DIV-002",
             "SDD §8.2 references equipment stowage area for 4-person crew. Bailout gas cylinders referenced.",
             "ASSUMED",
             "Equipment stowage layout to be confirmed in detailed design. Bailout gas capacity to be calculated."),
            ("8.3","Pt.5 Ch.10 Sec.8 [803]; DNV-OS-D201 Sec.1",
             "Umbilical management — Diver umbilical (gas, communications, hot water) to be managed to prevent entanglement. Surface umbilical (power, gas, comms) to be designed per DNV-OS-D201.",
             "SDD §8.3; RTM DIV-003; RTM SYS-003",
             "SDD §8.3 references umbilical management system with reel and guide. Hot water suit supply referenced. RTM SYS-003 records surface umbilical requirement and cross-references DNV-OS-D201 §1.",
             "ASSUMED",
             "Diver umbilical management system design to be submitted. Surface umbilical design to be confirmed against DNV-OS-D201 Sec.1 per RTM SYS-003. Separate umbilical calculations package required."),
            ("8.4","Pt.5 Ch.10 Sec.8 [804]",
             "Saturation schedule — Decompression tables and saturation schedule to be approved by competent authority.",
             "SDD §8.4; RTM DIV-004",
             "SDD §8.4 references approved saturation decompression schedule. Approval authority referenced.",
             "ASSUMED",
             "Approved saturation schedule to be included in operations manual. DNV confirmation of approval authority required."),
        ]
    },
    {
        "title": "Section 9 — Emergency Systems",
        "rows": [
            ("9.1","Pt.5 Ch.10 Sec.9 [901]",
             "Emergency shutdown — ESD system to isolate non-essential systems and maintain life-critical functions upon activation.",
             "SDD §9.1; RTM EMG-001",
             "SDD §9.1 references ESD system with manual and automatic activation. Life-critical systems maintained through ESD.",
             "ASSUMED",
             "ESD system design and logic to be submitted. ESD cause and effect matrix required."),
            ("9.2","Pt.5 Ch.10 Sec.9 [902]",
             "Emergency survival time — Habitat to sustain crew for minimum 72 hours on internal resources following loss of surface umbilical.",
             "SDD §9.2; RTM EMG-002; RTM BAT-002",
             "SDD §9.2 states emergency survival capability but endurance is contingent on battery capacity (40 kWh / 80 hr) which conflicts with RTM BAT-002 (60 kWh / 120 hr). Life support gas endurance not separately calculated in Rev 1.0.",
             "GAP",
             "Emergency survival time not demonstrably >= 72 hours. Battery conflict at Row 6.3 is a prerequisite. Life support gas endurance calculation absent. Action: Complete survival time analysis covering power, O2, CO2 scrubbing and water for 72-hour minimum. Link to battery conflict resolution."),
            ("9.3","Pt.5 Ch.10 Sec.9 [903]",
             "Hyperbaric Evacuation Vessel (HEV) — HEV or equivalent rescue means to be available for emergency transfer under pressure.",
             "SDD §9.3; RTM EMG-003",
             "SDD §9.3 references HEV requirement and states HEV to be provided by surface support vessel. HEV specification and compatibility with Sentinel 001 mating flange not confirmed in programme suite.",
             "TBD",
             "HEV specification and compatibility to be confirmed. Mating flange dimensions required (see Row 9.6). Owner to confirm HEV availability and compatibility with surface support vessel."),
            ("9.4","Pt.5 Ch.10 Sec.9 [904]",
             "Emergency breathing apparatus — EBA sets for all crew plus 25% spare. Accessible from all manned spaces.",
             "SDD §9.4; RTM EMG-004",
             "SDD §9.4 references EBA sets for 6 persons (crew 4 + 50% spare). Storage locations in all manned spaces.",
             "ASSUMED",
             "EBA set certification and capacity to be confirmed. Accessibility from all manned spaces to be verified."),
            ("9.5","Pt.5 Ch.10 Sec.9 [905]",
             "Emergency position indicating — EPIRB or acoustic pinger to be fitted. Activation automatic and manual.",
             "SDD §9.5; RTM EMG-005",
             "SDD §9.5 references acoustic pinger and emergency beacon with manual and automatic activation.",
             "ASSUMED",
             "Emergency position indicating equipment to be specified and included in commissioning tests."),
            ("9.6","Pt.5 Ch.10 Sec.9 [906]",
             "Rescue mating flange — Standard rescue mating flange to DNV or international standard to allow HEV connection.",
             "SDD §9.6; RTM EMG-006",
             "Rescue mating flange not addressed in SDD Rev 1.0. RTM EMG-006 records requirement. No flange specification, dimensions or design in programme suite.",
             "GAP",
             "Rescue mating flange absent from SDD Rev 1.0. Gap identified in Phase 2. Action: Specify and design rescue mating flange per applicable standard. Confirm compatibility with HEV (Row 9.3). Required before DNV plan approval."),
            ("9.7","Pt.5 Ch.10 Sec.9 [907]",
             "Emergency decompression — Procedures for emergency decompression of saturation divers in event of habitat breach.",
             "SDD §9.7; RTM EMG-007",
             "SDD §9.7 references emergency decompression procedures with approved table and compressed gas supply.",
             "ASSUMED",
             "Emergency decompression procedures to be included in operations manual. Gas supply calculation required."),
            ("9.8","Pt.5 Ch.10 Sec.9 [908]",
             "Diver recovery — Recovery procedures for incapacitated diver during saturation operations.",
             "SDD §9.8; RTM EMG-008",
             "Procedural requirement. SDD §9.8 references diver recovery procedures in operations manual. No specific hardware requirement for compliance map.",
             "N/A",
             "Procedural requirement addressed in operations manual. N/A for hardware compliance mapping — covered under operational procedures documentation."),
        ]
    },
    {
        "title": "Section 10 — Documentation & Certification",
        "rows": [
            ("10.1","Pt.5 Ch.10 Sec.10 [1001]",
             "DNV design approval — All drawings and calculations to be submitted to DNV for approval prior to fabrication.",
             "SDD §10.1; RTM DOC-001",
             "DNV design approval process initiated. AiP target November 2026. Full plan approval submission to follow AiP.",
             "ASSUMED",
             "Document submission schedule to DNV to be established. All gap and conflict items must be resolved before plan approval submission."),
            ("10.2","Pt.5 Ch.10 Sec.10 [1002]",
             "Calculations package — Structural, pressure, buoyancy, electrical and mechanical calculations as a consolidated package.",
             "SDD §10.2; RTM DOC-002",
             "Calculations package referenced in SDD §10.2. Individual calculations at various stages. Rows 3.3, 3.6 and 6.3 resolutions are prerequisites.",
             "ASSUMED",
             "Calculations package to be assembled and submitted. Prerequisites: conflict and gap resolutions at Rows 3.3, 3.6, 6.3."),
            ("10.3","Pt.5 Ch.10 Sec.10 [1003]",
             "Compliance matrix — DNV compliance map to be prepared and submitted demonstrating requirement traceability against all applicable governing standards.",
             "BT3345_9001_009 Rev 01 (this document)",
             "This document (BT3345_9001_009 Rev 01) constitutes the DNV Compliance Map for Sentinel 001. 78 requirements mapped across 12 sections covering DNV-RU-UWT Pt.5 Ch.10 (December 2024) as primary standard, and co-governing standards DNV-OS-E402 (July 2024), DNV-RU-OU-0375, DNV-OS-D201, DNV-OS-D202, and DNV-OS-B101 per Section 12.",
             "COMPLIANT",
             "This document. Rev 01 issued April 2026. Section 12 added to address co-governing standard gap identified post-initial issue. Next revision upon design updates or DNV feedback."),
            ("10.4","Pt.5 Ch.10 Sec.10 [1004]",
             "Test & commissioning plan — Documented plan for all factory acceptance tests, site acceptance tests and commissioning.",
             "SDD §10.3; RTM DOC-003",
             "Test and commissioning plan referenced in SDD §10.3 as planned document. Not yet issued.",
             "ASSUMED",
             "Test and commissioning plan to be developed and submitted prior to fabrication."),
            ("10.5","Pt.5 Ch.10 Sec.10 [1005]",
             "Operations manual — Comprehensive operations manual covering all normal and emergency procedures.",
             "SDD §10.4; RTM DOC-004",
             "Operations manual referenced in SDD §10.4 as planned document. Emergency procedures referenced at various SDD sections.",
             "ASSUMED",
             "Operations manual to be developed. Must incorporate CO2 setpoint corrections per Row 5.3 and approved saturation schedule."),
            ("10.6","Pt.5 Ch.10 Sec.10 [1006]",
             "Maintenance manual — Maintenance schedule and procedures for all systems and equipment.",
             "SDD §10.5; RTM DOC-005",
             "Maintenance manual referenced in SDD §10.5 as planned document.",
             "ASSUMED",
             "Maintenance manual to be developed. Must address battery maintenance given thermal runaway OI-06 (Row 6.5)."),
        ]
    },
    {
        "title": "Section 11 — Additional Requirements",
        "rows": [
            ("11.1","Pt.5 Ch.10 Sec.11 [1101]",
             "Moon-pool habitat specific — Additional requirements for open-bottom moon-pool habitat configurations.",
             "SDD §11.1; RTM SYS-010",
             "Sentinel 001 is an open-bottom moon-pool design. Specific requirements for this configuration to be confirmed with DNV. Applicability of enclosed-habitat requirements (e.g. bilge, Row 4.2) requires clarification.",
             "TBD",
             "DNV technical guidance on moon-pool specific requirements to be sought. Bilge applicability (Row 4.2) is primary outstanding question. Include in AiP pre-submission meeting agenda."),
        ]
    },
    {
        "title": "Section 12 — Co-Governing Standards (DNV-OS-E402, DNV-RU-OU-0375, DNV-OS-D201, DNV-OS-D202, DNV-OS-B101)",
        "rows": [
            ("12.1","Pt.5 Ch.10 Sec.2 [2.1]; DNV-OS-E402 (Jul 2024)",
             "Co-governing standard — LIVHAB(SAT) shall comply with DNV-OS-E402 (Diving Systems, July 2024). This is a mandatory co-requirement explicitly stated in Pt.5 Ch.10 Sec.2 [2.1], not a supplemental reference.",
             "Pt.5 Ch.10 Sec.2 [2.1]; Programme Context; SDD §1.0",
             "Pt.5 Ch.10 Sec.2 [2.1] states: 'Habitats that subject occupants to equivalent pressure at depth are considered to lead to saturation of occupants such LIVHAB(SAT) shall comply with the requirements of DNV-OS-E402 and DNV-RU-OU-0375 as applicable for habitats.' Programme Context confirms DNV-OS-E402 (July 2024) as supporting standard. DNV-OS-E402 compliance is mandatory for all SAT habitat compartments.",
             "ASSUMED",
             "DNV-OS-E402 compliance review not yet formally initiated as separate work package. Full DNV-OS-E402 requirement mapping required before plan approval. Programme team to engage DNV Høvik to confirm scope of DNV-OS-E402 applicability to Sentinel 001 (moon-pool, ambient-pressure habitat operating at saturation depth). Action owner: Larry Chase / Joe Reeves."),
            ("12.2","DNV-OS-E402 Ch.1 Sec.1 Table 4",
             "Area classification — Inner area and outer area designations per DNV-OS-E402 Table 4 to be applied to all Sentinel 001 compartments and used as basis for system design.",
             "SDD §2.3; RTM SYS-001; DNV-OS-E402 Ch.1 Sec.1 Table 4",
             "Per Pt.5 Ch.10 Sec.2 [2.2-2.3]: compartments at equivalent pressure to depth are 'inner area'; compartments at 1 atm are 'outer area' relative to inner compartments. Sentinel 001 primary habitable volume operates at 2.40 bar(a) (13.7 msw saturation pressure) — all manned compartments are inner area. No 1-atm compartments identified in SDD Rev 1.0.",
             "TBD",
             "Area classification schedule to be prepared per DNV-OS-E402 Ch.1 Sec.1 Table 4. All SDD and system design documents must use inner/outer area designations consistently. SDD-SENTINEL-001 and all deliverables to reference DNV-OS-E402 for these definitions. Confirm with DNV whether any 1-atm service areas exist in the design. Action owner: Larry Chase."),
            ("12.3","DNV-RU-OU-0375; Pt.5 Ch.10 Sec.2 [2.1]",
             "Co-governing standard — LIVHAB(SAT) shall comply with DNV-RU-OU-0375 (Diving Systems) as applicable. Mandatory co-requirement per Pt.5 Ch.10 Sec.2 [2.1].",
             "Pt.5 Ch.10 Sec.2 [2.1]; Programme Context",
             "DNV-RU-OU-0375 is cited alongside DNV-OS-E402 in Pt.5 Ch.10 Sec.2 [2.1] as a co-requirement for LIVHAB(SAT). Scope of applicability to Sentinel 001 (as a habitat vs. a standalone diving system) to be confirmed with DNV. DNV-RU-OU-0375 was not listed in the Programme Context supporting standards — potential gap in programme document suite.",
             "TBD",
             "DNV-RU-OU-0375 applicability to Sentinel 001 to be formally confirmed with DNV Høvik. If applicable, full requirement mapping required. Add DNV-RU-OU-0375 to programme document suite and SDD reference list. Action owner: Larry Chase / Oliver Hibbert (DNV point of contact). Include in AiP pre-submission agenda."),
            ("12.4","DNV-OS-D201 Sec.1; Pt.5 Ch.10 Sec.8 [803]",
             "Surface umbilical — Design of surface umbilical (power, gas, data, hot water from surface support vessel to habitat) to comply with DNV-OS-D201.",
             "SDD §6.1; SDD §8.3; RTM SYS-003; RTM ELE-001",
             "RTM SYS-003 records surface umbilical requirement with DNV-OS-D201 §1 as regulatory reference. SDD §6.1 documents surface umbilical power at 440V 3-phase, 22 kW normal load. SDD §8.3 references umbilical management system. DNV-OS-D201 governs design of umbilical systems for offshore units.",
             "ASSUMED",
             "Surface umbilical design calculation to be prepared and submitted per DNV-OS-D201 Sec.1. Umbilical package to include: power, gas, communications, hot water and emergency services. Umbilical termination and management system to be included. Confirm with DNV whether DNV-OS-D201 full edition (2024) applies. Action owner: Larry Chase / systems engineer."),
            ("12.5","DNV-OS-D202; RTM SYS-004",
             "Remote control and monitoring — Systems for remote operation, monitoring and control of habitat from surface to comply with DNV-OS-D202.",
             "SDD §7.1; SDD §7.5; RTM SYS-004; RTM MON-001",
             "RTM SYS-004 records real-time video, audio and sensor data requirement with DNV-OS-D202 (Remote Control) as regulatory reference. SDD §7.1 documents integrated environmental monitoring system (COMPLIANT — Row 7.1). SDD §7.5 references surface monitoring console.",
             "ASSUMED",
             "Remote control and monitoring system design to be confirmed against DNV-OS-D202 requirements. SAT test to be conducted per RTM SYS-004 verification method. Confirm with DNV whether DNV-OS-D202 full standard applies to Sentinel 001 habitat-to-surface interface. Action owner: Larry Chase / electrical/control engineer."),
            ("12.6","DNV-OS-B101 (2024); Pt.5 Ch.10 Sec.3 [302-303]",
             "Structural material standard — Structural steel and pressure vessel material to comply with DNV-OS-B101 (Metallic Materials) as supporting standard.",
             "SDD §3.0; SDD §3.2; RTM STR-001; Programme Context",
             "Programme Context lists DNV-OS-B101 (2024) as a supporting standard. DNV-OS-B101 governs metallic materials for offshore structures and is applicable to pressure hull steel selection. Current material conflict (Row 3.3): SDD specifies SA-516 Gr.70, RTM specifies A576 Gr.50. Neither has been formally assessed against DNV-OS-B101 material requirements.",
             "IN PROGRESS",
             "DNV-OS-B101 (2024) material assessment required for selected structural steel grade once Row 3.3 conflict is resolved. Material qualification plan (ESP-SAT-MAT-23-DNV) and material data record (MDR-SENTINEL-001-A) to be reviewed against DNV-OS-B101 requirements. DNV-OS-B101 assessment is BLOCKED pending Row 3.3 resolution. Action owner: Larry Chase / structural engineer / Joseph Welker."),
        ]
    },
]

# Verify row count
total = sum(len(s["rows"]) for s in SECTIONS)
print(f"Total rows: {total}")
assert total == 78, f"Expected 78 rows, got {total}"

# Verify confirmed status rows
STATUS_CHECK = {
    "1.5": "GAP", "2.4": "IN PROGRESS", "3.3": "CONFLICT", "3.6": "GAP",
    "3.12": "COMPLIANT", "4.2": "TBD", "5.3": "GAP",
    "6.1": "COMPLIANT", "6.2": "COMPLIANT", "6.3": "CONFLICT", "6.5": "IN PROGRESS",
    "7.1": "COMPLIANT", "7.3": "COMPLIANT", "9.2": "GAP", "9.3": "TBD",
    "9.6": "GAP", "10.3": "COMPLIANT",
    "11.1": "TBD", "12.2": "TBD", "12.3": "TBD", "12.6": "IN PROGRESS",
}
for sec in SECTIONS:
    for row in sec["rows"]:
        rn, _, _, _, _, st, _ = row
        if rn in STATUS_CHECK:
            assert st == STATUS_CHECK[rn], f"Row {rn}: expected {STATUS_CHECK[rn]}, got {st}"
print("Status checks passed.")

# Count totals
from collections import Counter
status_counts = Counter()
for sec in SECTIONS:
    for row in sec["rows"]:
        status_counts[row[5]] += 1
print("Status totals:", dict(status_counts))

# ─── Build the document ────────────────────────────────────────────────────────
doc = Document()

# Page setup: landscape US Letter, 0.75" margins
sec0 = doc.sections[0]
sec0.page_width  = Inches(11)
sec0.page_height = Inches(8.5)
sec0.orientation = WD_ORIENT.LANDSCAPE
sec0.top_margin    = Inches(0.75)
sec0.bottom_margin = Inches(0.75)
sec0.left_margin   = Inches(0.75)
sec0.right_margin  = Inches(0.75)

# Default style
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(9)

# ─── Header ───────────────────────────────────────────────────────────────────
hdr = sec0.header
hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = hp.add_run("BT3345_9001_009 Rev 01  |  Sentinel 001 DNV Compliance Map  |  DNV-RU-UWT Pt.5 Ch.10 (Dec 2024)  |  Co-governing: DNV-OS-E402 (Jul 2024), DNV-RU-OU-0375")
r1.font.name = "Arial"; r1.font.size = Pt(8); r1.font.color.rgb = NAVY_RGB
r1.font.bold = True
# Add right-aligned page number using tab
hp.paragraph_format.tab_stops.add_tab_stop(Inches(9.5), WD_ALIGN_PARAGRAPH.RIGHT)
r2 = hp.add_run("\t")
r2.font.name = "Arial"; r2.font.size = Pt(8)
add_run_field(hp, "PAGE", prefix="Page ", suffix=" of ", size=8)
add_run_field(hp, "NUMPAGES", size=8)

# ─── Footer ───────────────────────────────────────────────────────────────────
ftr = sec0.footer
fp = ftr.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
rf = fp.add_run("PROPRIETARY — Bastion Technologies, Inc.  |  Date: 2026-04-26  |  Classification: CONFIDENTIAL")
rf.font.name = "Arial"; rf.font.size = Pt(7); rf.font.color.rgb = NAVY_RGB

# ─── Title block ──────────────────────────────────────────────────────────────
def add_heading_para(doc, text, size=14, bold=True, color=NAVY_RGB, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = color
    return p

add_heading_para(doc, "DNV Compliance Map — Sentinel 001 Underwater Habitat", size=16, space_before=0)
add_heading_para(doc, "BT3345_9001_009 Rev 01  |  DNV-RU-UWT Pt.5 Ch.10 (Dec 2024)  +  DNV-OS-E402 (Jul 2024)  +  DNV-RU-OU-0375  |  Target: LIVHAB(SAT)", size=10, bold=False, color=NAVY_RGB, space_before=2, space_after=2)

# Meta-data table (4-col)
meta_tbl = doc.add_table(rows=3, cols=4)
meta_tbl.style = "Table Grid"
meta_data = [
    [("Document No.","BT3345_9001_009"),("Revision","Rev 01")],
    [("Project","Sentinel 001 Underwater Habitat — LIVHAB(SAT) AiP Programme"),("Date","2026-04-26")],
    [("Prepared by","Larry Chase, Senior Systems Engineer, Bastion Technologies, Inc."),("Approved by","Pending")],
]
for row_i, row_data in enumerate(meta_data):
    cells = meta_tbl.rows[row_i].cells
    col = 0
    for label, value in row_data:
        c1 = cells[col];   c2 = cells[col+1]
        shade_cell(c1, NAVY_HEX)
        shade_cell(c2, LBBLUE_HEX)
        set_cell_margins(c1); set_cell_margins(c2)
        cell_para(c1, label, bold=True, color=WHITE_RGB, size=8)
        cell_para(c2, value, size=8)
        col += 2
set_table_borders(meta_tbl)

# ASSUMED note (Correction 4)
note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(6)
note_p.paragraph_format.space_after  = Pt(2)
note_r1 = note_p.add_run("NOTE ON 'ASSUMED' STATUS: ")
note_r1.font.name = "Arial"; note_r1.font.size = Pt(8); note_r1.font.bold = True
note_r1.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)
note_r2 = note_p.add_run(
    "Rows marked ASSUMED indicate that compliance with the stated requirement is assumed based on declared design intent "
    "in SDD-SENTINEL-001 Rev 1.0 and RTM-SENT-001 Rev E, but has not yet been independently verified by calculation, "
    "test, or DNV review. ASSUMED rows require confirmation at plan approval stage. This document is issued at AiP "
    "programme level (target: November 2026).")
note_r2.font.name = "Arial"; note_r2.font.size = Pt(8)

# ─── Status legend ────────────────────────────────────────────────────────────
legend_p = add_heading_para(doc, "STATUS LEGEND", size=9, space_before=6, space_after=2)
legend_tbl = doc.add_table(rows=1, cols=7)
legend_tbl.style = "Table Grid"
legend_cells = legend_tbl.rows[0].cells
for i, (status, fmt) in enumerate(STATUS_FMT.items()):
    shade_cell(legend_cells[i], fmt["bg"])
    set_cell_margins(legend_cells[i], top=30, bottom=30, left=60, right=60)
    cell_para(legend_cells[i], status, bold=True, color=fmt["fg"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
set_table_borders(legend_tbl)

doc.add_paragraph().paragraph_format.space_before = Pt(4)

# ─── Main compliance table ────────────────────────────────────────────────────
# Columns: Row# | DNV Ref | Requirement Summary | Source Doc Ref | Compliance Evidence | Status | Notes/Actions
# Widths (inches): 0.4 | 1.0 | 2.0 | 1.1 | 2.2 | 0.75 | 2.05 = 9.5"
COL_WIDTHS = [0.4, 1.0, 2.0, 1.1, 2.2, 0.75, 2.05]
COL_HEADERS = ["Row #", "DNV Reference", "Requirement Summary", "Source Doc Ref", "Compliance Evidence", "Status", "Notes / Actions"]

comp_tbl = doc.add_table(rows=1, cols=7)
comp_tbl.style = "Table Grid"
set_table_borders(comp_tbl)

# Header row
hdr_cells = comp_tbl.rows[0].cells
for i, (htext, w) in enumerate(zip(COL_HEADERS, COL_WIDTHS)):
    shade_cell(hdr_cells[i], NAVY_HEX)
    set_cell_margins(hdr_cells[i])
    hdr_cells[i].width = Inches(w)
    cell_para(hdr_cells[i], htext, bold=True, color=WHITE_RGB, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

row_num_global = 0
for section_data in SECTIONS:
    # Section separator row
    sec_row = comp_tbl.add_row()
    sec_cells = sec_row.cells
    # Merge all 7 cells
    merged = sec_cells[0]
    for j in range(1, 7):
        merged = merged.merge(sec_cells[j])
    shade_cell(merged, LBBLUE_HEX)
    set_cell_margins(merged, top=30, bottom=30, left=80, right=80)
    cell_para(merged, section_data["title"].upper(), bold=True, color=NAVY_RGB, size=9)

    for row_tuple in section_data["rows"]:
        row_num, dnv_ref, req_sum, src_ref, evidence, status, notes = row_tuple
        row_num_global += 1
        fmt = STATUS_FMT[status]

        data_row = comp_tbl.add_row()
        cells = data_row.cells
        # Alternating row background (very light blue for odd data rows)
        alt_bg = "EEF4F8" if row_num_global % 2 == 0 else "FFFFFF"

        values = [row_num, dnv_ref, req_sum, src_ref, evidence, status, notes]
        for ci, (val, w) in enumerate(zip(values, COL_WIDTHS)):
            cells[ci].width = Inches(w)
            set_cell_margins(cells[ci])
            if ci == 5:  # Status column
                shade_cell(cells[ci], fmt["bg"])
                cell_para(cells[ci], val, bold=True, color=fmt["fg"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                shade_cell(cells[ci], alt_bg)
                is_bold = (ci == 0)
                txt_color = NAVY_RGB if ci == 0 else BLACK_RGB
                cell_para(cells[ci], val, bold=is_bold, color=txt_color, size=8)

# ─── Summary table ────────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading_para(doc, "COMPLIANCE SUMMARY", size=10, space_before=8, space_after=4)

sum_tbl = doc.add_table(rows=2, cols=7)
sum_tbl.style = "Table Grid"
set_table_borders(sum_tbl)
sum_hdr = sum_tbl.rows[0].cells
sum_val = sum_tbl.rows[1].cells
statuses = list(STATUS_FMT.keys())
for i, st in enumerate(statuses):
    fmt = STATUS_FMT[st]
    shade_cell(sum_hdr[i], fmt["bg"])
    set_cell_margins(sum_hdr[i], top=30, bottom=30, left=60, right=60)
    cell_para(sum_hdr[i], st, bold=True, color=fmt["fg"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(sum_val[i], LBBLUE_HEX)
    set_cell_margins(sum_val[i], top=30, bottom=30, left=60, right=60)
    cell_para(sum_val[i], str(status_counts.get(st, 0)), bold=True, color=NAVY_RGB, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total note
tot_p = doc.add_paragraph()
tot_p.paragraph_format.space_before = Pt(4)
tr = tot_p.add_run(f"Total requirements mapped: {total}  |  Primary standard: DNV-RU-UWT Pt.5 Ch.10 (December 2024)  |  COMPLIANT count confirmed: 6 rows (3.12, 6.1, 6.2, 7.1, 7.3, 10.3)")
tr.font.name = "Arial"; tr.font.size = Pt(8); tr.font.italic = True; tr.font.color.rgb = NAVY_RGB
tot_p2 = doc.add_paragraph()
tot_p2.paragraph_format.space_before = Pt(2)
tr2 = tot_p2.add_run(
    "Co-governing standards: DNV-OS-E402 (Jul 2024)  |  DNV-RU-OU-0375  |  DNV-OS-D201  |  DNV-OS-D202  |  DNV-OS-B101 (2024)  "
    "|  Section 12 added Rev 01 to address co-governing standard gap")
tr2.font.name = "Arial"; tr2.font.size = Pt(8); tr2.font.italic = True; tr2.font.color.rgb = NAVY_RGB

# ─── Conflict & Gap register summary ─────────────────────────────────────────
add_heading_para(doc, "BLOCKING CONFLICTS & OPEN GAPS REGISTER", size=10, space_before=10, space_after=4)

issues = [
    ("CONFLICT","3.3","Steel material specification","SDD: SA-516 Gr.70 vs RTM: A576 Gr.50","Larry Chase / structural engineer","SDD Rev 1.1 or RTM Rev F"),
    ("CONFLICT","6.3","Battery capacity & endurance","SDD: 40 kWh / 80 hr vs RTM: 60 kWh / 120 hr","Larry Chase / electrical engineer","SDD Rev 1.1 or RTM Rev F"),
    ("GAP","1.5","Safety philosophy / PHA document","PHA-SENTINEL-001 not in programme suite","Larry Chase","Issue PHA-SENTINEL-001"),
    ("GAP","3.6","Hinge fatigue analysis","Fatigue calculation absent from programme","Structural engineer","Commission analysis"),
    ("GAP","5.3","CO2 alarm pCO2 thresholds","SDD % vol limits exceed pCO2 limits at operating pressure","Larry Chase / Jim Williamson","Update SDD §4.4, RTM LSS-003, system config"),
    ("GAP","9.2","Emergency survival time calculation","72-hr endurance not demonstrated","Larry Chase","Complete survival time analysis"),
    ("GAP","9.6","Rescue mating flange","Not specified in SDD Rev 1.0","Larry Chase","Specify and design per standard"),
]
issue_tbl = doc.add_table(rows=1, cols=6)
issue_tbl.style = "Table Grid"
set_table_borders(issue_tbl)
issue_hdrs = ["Type","Row","Subject","Issue Description","Action Owner","Resolution Required"]
issue_widths = [0.7, 0.4, 1.4, 2.4, 1.6, 2.0]
ih_cells = issue_tbl.rows[0].cells
for i, (h, w) in enumerate(zip(issue_hdrs, issue_widths)):
    shade_cell(ih_cells[i], NAVY_HEX)
    ih_cells[i].width = Inches(w)
    set_cell_margins(ih_cells[i])
    cell_para(ih_cells[i], h, bold=True, color=WHITE_RGB, size=8)

for idx, (itype, row_n, subj, desc, owner, res) in enumerate(issues):
    fmt = STATUS_FMT[itype]
    alt = "FFFFFF" if idx % 2 == 0 else "EEF4F8"
    ir = issue_tbl.add_row()
    ic = ir.cells
    vals = [itype, row_n, subj, desc, owner, res]
    for ci, (v, w) in enumerate(zip(vals, issue_widths)):
        ic[ci].width = Inches(w)
        set_cell_margins(ic[ci])
        if ci == 0:
            shade_cell(ic[ci], fmt["bg"])
            cell_para(ic[ci], v, bold=True, color=fmt["fg"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            shade_cell(ic[ci], alt)
            cell_para(ic[ci], v, size=8)

# ─── Closing note ─────────────────────────────────────────────────────────────
doc.add_paragraph()
close_p = doc.add_paragraph()
close_p.paragraph_format.space_before = Pt(6)
cr = close_p.add_run(
    "This compliance map is issued at Approval in Principle (AiP) programme stage. All CONFLICT and GAP items must be "
    "resolved prior to DNV plan approval submission. ASSUMED rows require verification by calculation, test or DNV survey "
    "before COMPLIANT status can be confirmed. Section 12 (rows 12.1-12.6) documents co-governing standard obligations "
    "under DNV-OS-E402 (July 2024), DNV-RU-OU-0375, DNV-OS-D201, DNV-OS-D202 and DNV-OS-B101 — these were not mapped "
    "in the initial issue and require formal compliance review. DNV-OS-E402 and DNV-RU-OU-0375 are mandatory per "
    "Pt.5 Ch.10 Sec.2 [2.1]. RTM cover page revision discrepancy (title cell: Rev D | 13 Feb 2026 vs "
    "revision detail: Rev E | 2026-03-20) is flagged as a document integrity issue — action owner: Larry Chase.")
cr.font.name = "Arial"; cr.font.size = Pt(8); cr.font.italic = True

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\lchas\DNV _System\Sentinel001_DNV_Compliance_Map_Rev01.docx"
doc.save(out_path)
print(f"\nDocument saved: {out_path}")
print("Phase 3 complete.")


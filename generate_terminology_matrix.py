#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate BT3345_9001_008 - DNV Terminology and Requirements Matrix
Sentinel 001 Programme | Bastion Technologies, Inc.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── CONSTANTS ─────────────────────────────────────────────────────────────────
DOC_NUMBER      = "BT3345_9001_008"
DOC_TITLE       = "DNV Terminology and Requirements Matrix"
DOC_SUBTITLE    = "Source: DNV-RU-UWT Pt.5 Ch.10 (December 2024)"
DOC_REV         = "00"
DOC_DATE        = "24 April 2026"
DOC_STATUS      = "DRAFT - FOR INTERNAL REVIEW"
DOC_OWNER       = "Bastion Technologies, Inc."
DOC_PROGRAMME   = "Sentinel 001"
DOC_PREPARED_BY = "[INSERT NAME]"
DOC_APPROVED_BY = "[INSERT NAME]"
OUT_PATH        = r"C:\Users\lchas\DNV _System\Sentinel001_DNV_Terminology_Matrix_Rev00.docx"

NAVY        = "1F3864"
LIGHT_BLUE  = "D5E8F0"
WHITE       = "FFFFFF"
LIGHT_RED   = "FFE0E0"
LIGHT_YELLOW= "FFF2CC"
LIGHT_GREEN = "E2EFDA"
BORDER      = "CCCCCC"
NAVY_RGB    = RGBColor(0x1F, 0x38, 0x64)
WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)

# ── XML HELPERS ───────────────────────────────────────────────────────────────

def set_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    tcPr.append(shd)


def set_borders(cell, color=BORDER, hide=False):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        if hide:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color.upper())
        tcB.append(el)
    tcPr.append(tcB)


def set_tbl_width(tbl, inches):
    tbl_el = tbl._tbl
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(inches * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def set_para_spacing(para, before=0, after=40):
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:spacing")):
        pPr.remove(existing)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def cell_write(cell, text, bold=False, color_rgb=None, size=9,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_para_spacing(p, before=0, after=40)
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb


def add_field_run(para, field_name, size_pt=8):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(size_pt * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(rFonts)
    rPr.append(sz_el)
    rPr.append(szCs)
    run = OxmlElement("w:r")
    run.append(rPr)
    run.append(fld_begin)
    run.append(instr)
    run.append(fld_end)
    para._p.append(run)


def add_page_xy(para, size_pt=8):
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _txt(t):
        r = para.add_run(t)
        r.font.name = "Arial"
        r.font.size = Pt(size_pt)

    _txt("Page ")
    add_field_run(para, "PAGE", size_pt)
    _txt(" of ")
    add_field_run(para, "NUMPAGES", size_pt)


# ── PAGE SETUP ────────────────────────────────────────────────────────────────

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width  = Inches(11)
    sec.page_height = Inches(8.5)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sectPr = sec._sectPr
    pgSz = sectPr.find(qn("w:pgSz"))
    if pgSz is None:
        pgSz = OxmlElement("w:pgSz")
        sectPr.append(pgSz)
    pgSz.set(qn("w:orient"), "landscape")


def setup_hf(doc):
    sec = doc.sections[0]
    usable = 9.5

    # ── FOOTER ──────────────────────────────────────────────────────────
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    ft = footer.add_table(1, 3, Inches(usable))
    ft.style = "Table Grid"
    set_tbl_width(ft, usable)
    for row in ft.rows:
        for cell in row.cells:
            set_borders(cell, hide=True)

    ft.cell(0, 0).width = Inches(3.0)
    ft.cell(0, 1).width = Inches(3.5)
    ft.cell(0, 2).width = Inches(3.0)

    lc = ft.cell(0, 0)
    lc.text = ""
    lp = lc.paragraphs[0]
    set_para_spacing(lp)
    r = lp.add_run(f"{DOC_NUMBER} Rev {DOC_REV}  |  DRAFT")
    r.font.name = "Arial"; r.font.size = Pt(8)

    cc = ft.cell(0, 1)
    cc.text = ""
    cp = cc.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cp)
    r2 = cp.add_run("SENTINEL 001  |  BASTION TECHNOLOGIES, INC.")
    r2.font.name = "Arial"; r2.font.size = Pt(8)

    rc = ft.cell(0, 2)
    rc.text = ""
    rp = rc.paragraphs[0]
    set_para_spacing(rp)
    add_page_xy(rp)

    # ── HEADER ──────────────────────────────────────────────────────────
    header = sec.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    ht = header.add_table(1, 2, Inches(usable))
    ht.style = "Table Grid"
    set_tbl_width(ht, usable)
    for row in ht.rows:
        for cell in row.cells:
            set_borders(cell, hide=True)

    ht.cell(0, 0).width = Inches(5.5)
    ht.cell(0, 1).width = Inches(4.0)

    lh = ht.cell(0, 0)
    lh.text = ""
    lhp = lh.paragraphs[0]
    set_para_spacing(lhp, before=0, after=20)
    r3 = lhp.add_run(DOC_TITLE)
    r3.font.name = "Arial"; r3.font.size = Pt(8); r3.font.italic = True

    rh = ht.cell(0, 1)
    rh.text = ""
    rhp = rh.paragraphs[0]
    rhp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(rhp, before=0, after=20)
    r4 = rhp.add_run("DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    r4.font.name = "Arial"; r4.font.size = Pt(8); r4.font.italic = True


# ── CONTENT HELPERS ───────────────────────────────────────────────────────────

def add_title_bar(doc):
    tbl = doc.add_table(1, 1)
    tbl.style = "Table Grid"
    set_tbl_width(tbl, 9.5)
    cell = tbl.cell(0, 0)
    set_shading(cell, NAVY)
    set_borders(cell, color=NAVY)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=80, after=80)
    r1 = p.add_run(DOC_TITLE + "\n")
    r1.font.name = "Arial"; r1.font.size = Pt(16)
    r1.font.bold = True; r1.font.color.rgb = WHITE_RGB
    r2 = p.add_run(DOC_SUBTITLE)
    r2.font.name = "Arial"; r2.font.size = Pt(11)
    r2.font.bold = False; r2.font.color.rgb = WHITE_RGB


def add_header_block(doc):
    fields = [
        ("Document Number",   DOC_NUMBER),
        ("Title",             DOC_TITLE),
        ("Subtitle",          DOC_SUBTITLE),
        ("Revision",          DOC_REV),
        ("Date",              DOC_DATE),
        ("Status",            DOC_STATUS),
        ("Owner",             DOC_OWNER),
        ("Programme",         DOC_PROGRAMME),
        ("Prepared by",       DOC_PREPARED_BY),
        ("Approved by",       DOC_APPROVED_BY),
    ]
    doc.add_paragraph()
    tbl = doc.add_table(len(fields), 2)
    tbl.style = "Table Grid"
    set_tbl_width(tbl, 6.0)
    for i, (label, value) in enumerate(fields):
        row = tbl.rows[i]
        bg = LIGHT_BLUE if i % 2 == 1 else WHITE
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = Inches(1.5)
        vc.width = Inches(4.5)
        set_shading(lc, NAVY)
        set_borders(lc, BORDER)
        set_shading(vc, bg)
        set_borders(vc, BORDER)
        cell_write(lc, label, bold=True, color_rgb=WHITE_RGB, size=9)
        cell_write(vc, value, size=9)


def add_section_heading(doc, text, before=160, after=60):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(12)
    run.font.bold = True; run.font.color.rgb = NAVY_RGB


def add_source_note(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=60)
    run = p.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(9); run.font.italic = True


def add_body_para(doc, text, before=0, after=80):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after)
    run = p.add_run(text)
    run.font.name = "Arial"; run.font.size = Pt(9)


def make_table(doc, headers, data_rows, col_widths, font_size=9):
    n_cols = len(headers)
    n_rows = len(data_rows)
    tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_tbl_width(tbl, sum(col_widths))

    # Apply column widths to every row
    for row in tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

    # Header row
    hdr_row = tbl.rows[0]
    for i, hdr in enumerate(headers):
        c = hdr_row.cells[i]
        set_shading(c, NAVY)
        set_borders(c, BORDER)
        cell_write(c, hdr, bold=True, color_rgb=WHITE_RGB, size=font_size)

    # Data rows
    for ri, row_data in enumerate(data_rows):
        bg = LIGHT_BLUE if ri % 2 == 1 else WHITE
        tbl_row = tbl.rows[ri + 1]
        for ci, txt in enumerate(row_data):
            c = tbl_row.cells[ci]
            set_shading(c, bg)
            set_borders(c, BORDER)
            cell_write(c, str(txt), size=font_size)

    return tbl


def add_flag_box(doc, flag_num, severity, title, body, bg_hex, before=80):
    p_spacer = doc.add_paragraph()
    set_para_spacing(p_spacer, before=before, after=0)

    tbl = doc.add_table(1, 1)
    tbl.style = "Table Grid"
    set_tbl_width(tbl, 9.5)

    cell = tbl.cell(0, 0)
    set_shading(cell, bg_hex)
    set_borders(cell, BORDER)
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p1, before=60, after=40)
    r1 = p1.add_run(f"FLAG {flag_num}  |  {severity}  |  {title}")
    r1.font.name = "Arial"; r1.font.size = Pt(9); r1.font.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p2, before=0, after=60)
    r2 = p2.add_run(body)
    r2.font.name = "Arial"; r2.font.size = Pt(9)


def add_doc_control(doc):
    add_section_heading(doc, "DOCUMENT CONTROL", before=200)
    headers = ["Rev", "Date", "Description", "Author", "Approved"]
    data = [
        ["00", DOC_DATE, "Initial issue - DNV Terminology Matrix, Rev 00", "[INSERT NAME]", "[INSERT NAME]"],
    ]
    col_widths = [0.5, 1.2, 4.5, 1.7, 1.6]
    make_table(doc, headers, data, col_widths, font_size=9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_para_spacing(p, before=80, after=40)
    r = p.add_run("PROPRIETARY NOTICE")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=40)
    r2 = p2.add_run(
        "This document contains proprietary information belonging to Bastion Technologies, Inc. "
        "It is submitted in confidence and shall not be reproduced, copied, disclosed to a third party, "
        "or used for any purpose other than that for which it is supplied, without the express written "
        "permission of Bastion Technologies, Inc. The recipient of this document agrees to hold its "
        "contents in confidence. This document is issued for the sole purpose of the Sentinel 001 "
        "AiP programme with DNV Hovik."
    )
    r2.font.name = "Arial"; r2.font.size = Pt(8); r2.font.italic = True


# ── TABLE DATA ────────────────────────────────────────────────────────────────

PART_A_HEADERS = ["ID", "Term", "Type", "DNV Definition", "Source Clause",
                  "Sentinel 001 Application", "Compliance Flag", "Confidence"]
PART_A_WIDTHS  = [0.40, 0.80, 0.70, 2.40, 0.90, 2.30, 1.30, 0.70]

PART_A_DATA = [
    ["A-01", "LIVHAB", "Class Notation",
     "Class notation for live in underwater habitats. Provides technical and procedural framework "
     "supporting design, fabrication, installation, and testing of a habitat system.",
     "Sec.1 [2], [6]",
     "GOVERNING NOTATION for Sentinel 001.",
     "N/A", "HIGH"],

    ["A-02", "ATMOS", "Notation Qualifier",
     "Operating condition where internal pressure is equal to approximately that of surface "
     "pressure (1 atmosphere).",
     "Sec.1 [6], Table 1",
     "NOT APPLICABLE to Sentinel 001 primary operating mode.",
     "Design basis is SAT -- do not conflate ATMOS atmospheric limits with SAT limits.",
     "HIGH"],

    ["A-03", "SAT", "Notation Qualifier",
     "Operating condition where typical internal pressure is equal to that of depth of operation "
     "of habitat. Leads to saturation of occupants.",
     "Sec.1 [6], Table 1; Sec.2 [2.1]",
     "GOVERNING QUALIFIER -- Sentinel 001 operates at 2.40 bar(a) / 13.7 msw. "
     "All SAT sub-clauses throughout document are mandatory.",
     "All SAT-specific clauses apply. DNV-OS-E402 and DNV-RU-OU-0375 are co-requirements.",
     "HIGH"],

    ["A-04", "LIVHAB(SAT)", "Programme Classification",
     "Combined notation: habitat operating at ambient pressure equal to depth. Must comply with "
     "DNV-OS-E402 and DNV-RU-OU-0375 as applicable for habitats.",
     "Sec.2 [2.1]",
     "Confirmed Sentinel 001 classification per programme context.",
     "DNV-OS-E402 compliance is mandatory -- not supplemental.",
     "HIGH"],
]

PART_B_HEADERS = ["ID", "Term", "DNV Definition (verbatim/paraphrased)", "Source Clause",
                  "Sentinel 001 Application", "Compliance Flag", "Confidence"]
PART_B_WIDTHS  = [0.40, 1.00, 2.40, 0.80, 2.50, 1.70, 0.70]

PART_B_DATA = [
    ["B-01", "ambient pressure habitat",
     "Typical internal pressure is equal to that of depth of operation of habitat.",
     "Sec.1 [8.1]",
     "Sentinel 001 nominal operating mode: 2.40 bar(a) at moon pool.",
     "Same as SAT mode.", "HIGH"],

    ["B-02", "autonomous habitat",
     "Habitat not physically connected to supporting infrastructure during operation (e.g. by umbilical).",
     "Sec.1 [8.1]",
     "TBD -- CONOPS must specify. Programme references umbilical connection (Pt.4 Ch.5).",
     "Review CONOPS -- determine autonomous vs. non-autonomous classification.", "MEDIUM"],

    ["B-03", "central control position",
     "Area designated for containing control and monitoring instrumentation of habitat life support systems.",
     "Sec.1 [8.1]",
     "Required design element in Sentinel 001 -- ref SDD-SENTINEL-001.",
     "Must be identifiable on GA plan submitted to DNV.", "HIGH"],

    ["B-04", "chamber",
     "Decompression, pressure or compression chambers; pressure vessels for human occupancy.",
     "Sec.1 [8.1]",
     "Any decompression chambers in Sentinel 001 -- Ch.2 compliance if present.",
     "Sec.3 [3] governs habitat decompression chamber requirements.", "HIGH"],

    ["B-05", "compartment",
     "Part(s) of a hull sufficiently large to contain at least one person; may have internal "
     "pressure or atmosphere different from adjacent compartments.",
     "Sec.1 [8.1]",
     "Sentinel 001: 120 m3 habitable volume, 9 persons. Compartment sizing drives hatch, "
     "trunk, and fire zone requirements.",
     "Minimum hatch clear diameter 600 mm per Sec.3 [5.8].", "HIGH"],

    ["B-06", "exostructure",
     "External cladding, supporting structures and fixtures outside the hull (all structure "
     "external to main hull).",
     "Sec.1 [8.1]",
     "Sentinel 001 exostructure -- structural lead Joseph Welker.",
     "OI-03 RELEVANT: escape hatch hinge is exostructure/hull interface; FLS analysis required.", "HIGH"],

    ["B-07", "essential services / equipment",
     "Services needing continuous operation to maintain habitat system's functionality re: "
     "safety, health, and life-sustaining environment.",
     "Sec.1 [8.1]",
     "Must be enumerated and protected in Sentinel 001 FMEA and emergency power design.",
     "LINKED TO OI-07: FMEA structural tab 3/10 -- essential services not adequately characterized.", "HIGH"],

    ["B-08", "habitat",
     "Live in underwater habitat (shorthand used throughout document).",
     "Sec.1 [8.1]",
     "Sentinel 001 = 'habitat' throughout all DNV submission documents.",
     "Terminology consistency required across entire programme document suite.", "HIGH"],

    ["B-09", "habitat decompression chamber",
     "Designated independently controlled pressure compartment part of habitat system, with "
     "facility to compress/decompress occupants.",
     "Sec.1 [8.1]",
     "If Sentinel 001 uses habitat itself as decompression chamber: full equivalence analysis "
     "required per Sec.3 [3.3].",
     "Verify in CONOPS -- additional Ch.2 requirements apply if applicable.", "MEDIUM"],

    ["B-10", "hull",
     "Watertight boundary of a habitat.",
     "Sec.1 [8.1]",
     "Sentinel 001 primary pressure hull -- ASTM A576 Grade 50 (345 MPa yield / 450 MPa UTS / "
     "167 MPa allowable).",
     "Structural design governing boundary. Safety factors per Sec.3 [1.7] apply.", "HIGH"],

    ["B-11", "independent habitat",
     "Habitat able to operate offshore and return occupants to place of safety without a support vessel.",
     "Sec.1 [8.1]",
     "TBD -- dependent on rescue concept design.",
     "If non-independent: support vessel/surface asset must be within class scope per Sec.2 [1.5].", "MEDIUM"],

    ["B-12", "live in underwater habitat",
     "Fully or partially submerged structure, typically comprising hull, pressure hull, "
     "compartments, chambers, and exostructure intended for human habitation.",
     "Sec.1 [8.1]",
     "GOVERNING DEFINITION of Sentinel 001 as a classified object.",
     "All programme documents must cite this definition as the regulatory baseline.", "HIGH"],

    ["B-13", "life support systems",
     "All systems needed to support life in a normal and emergency situation.",
     "Sec.1 [8.1]",
     "Sentinel 001 LSS -- Jim Williamson (Unique Group) lead. Full technical scope in Sec.5.",
     "Section 5 is the governing technical requirements section.", "HIGH"],

    ["B-14", "manned underwater operations",
     "When humans stay below surface or are exposed to increased ambient pressure.",
     "Sec.1 [8.1]",
     "Primary Sentinel 001 operation at 13.7 msw / 2.40 bar(a).",
     "Triggers all applicable regulatory requirements under DNV-RU-UWT.", "HIGH"],

    ["B-15", "non-autonomous habitat",
     "Habitat physically connected to supporting infrastructure (e.g. umbilical) during operation.",
     "Sec.1 [8.1]",
     "Likely applies to Sentinel 001 if umbilical-connected -- all support infrastructure "
     "brought within class scope.",
     "Confirm in CONOPS. Sec.2 [1.5] requirement applies.", "MEDIUM"],

    ["B-16", "normal operational period",
     "Period of time in which the habitat is planned to be in operation.",
     "Sec.1 [8.1]",
     "Must be explicitly defined in CONOPS (Z100 doc) -- drives all survival time and "
     "endurance calculations.",
     "120-hr battery endurance is the emergency basis. Normal operational period must be "
     "separately defined and is a distinct parameter.", "HIGH"],

    ["B-17", "place of safety",
     "Safe onshore or offshore location/vessel to which persons can be transferred in emergency "
     "where control of decompression is not considered life-threatening.",
     "Sec.1 [8.1]",
     "Critical definition for Sentinel 001 rescue concept. Must be identified and accessible "
     "within the rescue timeline.",
     "Decompression management is a requirement of the 'place of safety' definition. Cannot be "
     "a surface vessel without recompression capability.", "HIGH"],

    ["B-18", "pressure hull",
     "Hull which can withstand external pressure.",
     "Sec.1 [8.1]",
     "Sentinel 001 structural shell -- ASTM A576 Grade 50 steel. Safety factors per Sec.3 [1.7] apply.",
     "S2 = CDP/NDP >= 2.0; S1 = TDP/NDP >= 1.25.", "HIGH"],

    ["B-19", "support module",
     "Discrete machinery/equipment providing service to habitat (e.g. battery module, gas bank, "
     "windmill, surface buoy).",
     "Sec.1 [8.1]",
     "Sentinel 001 60 kWh battery system (2x30 kWh LiFePO4) is a support module.",
     "Per Sec.2 [1.5]: support assets must be within class scope or separately classed.", "HIGH"],

    ["B-20", "1-atmosphere habitat",
     "Internal pressure approximately equal to surface pressure.",
     "Sec.1 [8.1]",
     "NOT the Sentinel 001 operating mode. Provided to distinguish from SAT mode.",
     "Atmospheric limit Sec.5 [2.1] (0.19-0.23 bar ppO2) applies to ATMOS only. "
     "Sec.5 [3] governs Sentinel 001.", "HIGH"],

    ["B-21", "working devices",
     "Devices attachable to habitat for auxiliary purpose (e.g. manipulator arm, lifting device).",
     "Sec.1 [8.1]",
     "Scope per Sentinel 001 CONOPS. Any working devices require risk assessment per Sec.2 [3.5].",
     "Scientific testing activities listed in appendix to class.", "MEDIUM"],
]

PART_C_HEADERS = ["ID", "Abbreviation", "Full Term", "Source Clause",
                  "Sentinel 001 Relevance", "Programme Priority", "Confidence"]
PART_C_WIDTHS  = [0.40, 0.90, 1.30, 0.80, 2.80, 2.60, 0.70]

PART_C_DATA = [
    ["C-01", "ALS", "Accidental limit state", "Sec.1 [8.2]",
     "One of three structural limit states for Sentinel 001 hull analysis.",
     "Structural design.", "HIGH"],

    ["C-02", "AMV", "Average minute volume", "Sec.1 [8.2]",
     "LSS sizing parameter -- 9 persons at 13.7 msw; drives O2 demand and CO2 scrubber sizing.",
     "Jim Williamson (Unique Group).", "HIGH"],

    ["C-03", "CONOPS", "Concept of operations statement", "Sec.1 [8.2]",
     "Required document (Z100, FI); forms basis for all Sentinel 001 design decisions.",
     "AiP prerequisite.", "HIGH"],

    ["C-04", "EES", "Electrical energy storage system", "Sec.1 [8.2]",
     "Sentinel 001 EES: 60 kWh total (2x30 kWh LiFePO4). Governs emergency power calculations.",
     "Oguzhan Kilic (Electrical). See also Flag 7.", "HIGH"],

    ["C-05", "FMEA", "Failure mode and effect analysis", "Sec.1 [8.2]",
     "OI-07 open item -- structural FMEA tab scored 3/10 -- not AiP-ready.",
     "OI-07 -- Ken Ngo. See Flag 6.", "HIGH"],

    ["C-06", "FLS", "Fatigue limit state", "Sec.1 [8.2]",
     "Directly relevant to OI-03 -- escape hatch hinge fatigue analysis; RPN 175.",
     "OI-03 -- Joseph Welker.", "HIGH"],

    ["C-07", "MAWP", "Maximum allowable working pressure", "Sec.1 [8.2]",
     "S4 = PP/MAWP = 1.5 minimum (internal pressure test requirement).",
     "Structural test basis.", "HIGH"],

    ["C-08", "SWL", "Safe working load for lifting", "Sec.1 [8.2]",
     "Lifting points, exostructure fixation per Sec.2 [12.2]. Rupture load of anchor chains "
     "and wire ropes is governing design load.",
     "Structural design parameter.", "HIGH"],

    ["C-09", "ULS", "Ultimate limit state", "Sec.1 [8.2]",
     "Governing structural limit state for Sentinel 001 pressure hull.",
     "Structural design basis.", "HIGH"],
]

PART_D_HEADERS = ["ID", "Term", "Definition", "Source Clause",
                  "Safety Factor", "Sentinel 001 Application", "Confidence"]
PART_D_WIDTHS  = [0.40, 0.80, 1.50, 1.00, 1.00, 4.10, 0.70]

PART_D_DATA = [
    ["D-01", "NDP", "Nominal diving pressure", "Sec.3 [1.7], Table 1",
     "Basis",
     "2.40 bar(a) at 13.7 msw (moon pool operating depth). Governs hull design pressure basis.", "HIGH"],

    ["D-02", "TDP", "Test diving pressure", "Sec.3 [1.7], Table 1",
     "S1 = TDP/NDP >= 1.25",
     "TDP >= 3.00 bar(a). External pressure test requirement for hull qualification.", "HIGH"],

    ["D-03", "CDP", "Collapse diving pressure", "Sec.3 [1.7], Table 1",
     "S2 = CDP/NDP >= 2.0",
     "CDP >= 4.80 bar(a). Hull collapse resistance requirement. "
     "Structural adequacy basis for ASTM A576 Gr50 shell.", "HIGH"],

    ["D-04", "Deployment diving pressure (external)", "Deployment pressure -- external (S3)", "Sec.3 [1.7], Table 1",
     "S3 = defined by design risk assessment",
     "Not a fixed factor. Must be derived and justified by FMEA output. "
     "Critical for any deployment scenarios beyond 13.7 msw.", "HIGH"],

    ["D-05", "PB (MAWP)", "Maximum allowable working pressure (internal)", "Sec.3 [1.7], Table 2",
     "Basis",
     "Internal pressure design limit. Must account for max excursion: 2.99 bar(a) at 19.8 msw / 65 ft.", "HIGH"],

    ["D-06", "PP", "Test pressure (internal)", "Sec.3 [1.7], Table 2",
     "S4 = PP/PB >= 1.5",
     "PP >= 1.5 x PB. Pressure test requirement for internal systems and chambers.", "HIGH"],

    ["D-07", "inner area", "Compartment at equivalent pressure to depth -- SAT designation",
     "Sec.2 [2.2]; DNV-OS-E402 Ch.1 Sec.1 Table 4",
     "N/A",
     "All habitable Sentinel 001 compartments at operating depth are inner area. "
     "NOTE: term defined in DNV-OS-E402, not in Pt.5 Ch.10 directly. Cross-reference required.", "HIGH"],

    ["D-08", "outer area", "Compartment at 1 atmosphere relative to inner area -- SAT designation",
     "Sec.2 [2.3]; DNV-OS-E402 Ch.1 Sec.1 Table 4",
     "N/A",
     "Any 1-atm service areas in Sentinel 001 if present. "
     "NOTE: term defined in DNV-OS-E402, not in Pt.5 Ch.10 directly. Cross-reference required.", "HIGH"],
]

PART_E_HEADERS = ["ID", "Parameter", "Limit / Value", "Source Clause",
                  "Controlling for Sentinel 001?", "Programme Value",
                  "Gap / Flag", "Confidence"]
PART_E_WIDTHS  = [0.40, 1.30, 1.30, 0.80, 0.90, 1.20, 2.90, 0.70]

PART_E_DATA = [
    ["E-01", "ppO2 -- ATMOS normal range",
     "0.19 bar (min) to 0.23 bar (max)",
     "Sec.5 [2.1]", "NO -- superseded by SAT clause for Sentinel 001",
     "Nominal ppO2 = 0.50 bar (breathing air at 13.7 msw)",
     "Do NOT apply this limit to Sentinel 001 design. SAT clause Sec.5 [3] governs.", "HIGH"],

    ["E-02", "ppO2 -- SAT operating range",
     "0.19 bar (min) to 1.6 bar (max, short-term)",
     "Sec.5 [3]", "YES -- governing ppO2 limit for Sentinel 001",
     "Nominal 0.50 bar; max excursion 0.63 bar at 19.8 msw",
     "Programme target 0.4-0.6 bar: CONFORMS. Cite Sec.5 [3] in all DNV submissions.", "HIGH"],

    ["E-03", "ppO2 -- damage condition",
     "0.19 bar (min) to 1.6 bar (max, short-term, per diving tables)",
     "Sec.5 [2.1]", "YES",
     "Max 0.63 bar at 19.8 msw",
     "CONFORMS. Reference USN Diving Manual depth tables.", "HIGH"],

    ["E-04", "pCO2 -- normal limit",
     "< 0.005 bar (permanently maintained)",
     "Sec.5 [2.2]", "YES",
     "Not specified in programme context",
     "GAP: CO2 scrubber design must target < 0.005 bar. Jim Williamson to verify sizing for 9 persons.", "HIGH"],

    ["E-05", "pCO2 -- emergency limit",
     "< 0.02 bar (maintained for minimum survival time)",
     "Sec.5 [2.3]", "YES",
     "Not specified in programme context",
     "Emergency CO2 scrubber capacity must sustain 9 persons for full survival time. Verify sizing.", "HIGH"],

    ["E-06", "Standard man -- O2 consumption",
     "26 l/hr at 1 atm (average)",
     "Sec.5 [1.5], Table 1", "YES",
     "9 persons = 234 l/hr O2 baseline",
     "LSS design sizing input. Working occupants have higher demand. Use conservative load case.", "HIGH"],

    ["E-07", "Standard man -- CO2 production",
     "22 l/hr at 1 atm (average); 50 l/hr (working/divers)",
     "Sec.5 [1.5], Table 1", "YES",
     "9 persons avg: 198 l/hr; working: 450 l/hr",
     "CO2 scrubber sizing must use working-occupant value (50 l/hr/person) for conservative design.", "HIGH"],

    ["E-08", "RQ (respiration quotient)",
     "0.85 (volumetric CO2 produced / O2 consumed)",
     "Sec.5 [1.5], Table 1", "YES",
     "Applied in LSS calculations",
     "LSS design calculation input. Required for accurate gas balance analysis.", "HIGH"],
]

PART_F_HEADERS = ["ID", "Term", "Definition / Requirement", "Source Clause",
                  "Sentinel 001 Application", "Owner", "Confidence"]
PART_F_WIDTHS  = [0.40, 0.90, 2.40, 0.80, 2.90, 1.40, 0.70]

PART_F_DATA = [
    ["F-01", "fire zone",
     "Minimum 2 fire zones required. No zone to exceed 30 m2 floor area. "
     "Boundaries minimum class A60.",
     "Sec.8 [2.1]",
     "Sentinel 001 must be divided into minimum 2 fire zones -- required on GA plan submitted to DNV.",
     "Chrystal Revis (Fire/Safety)", "HIGH"],

    ["F-02", "A60",
     "Fire division class -- boundaries must withstand fire and pressure effects from extinguishers. "
     "Minimum standard for structural fire boundaries.",
     "Sec.8 [2.1]",
     "A60 is the minimum fire division standard for all structural fire boundaries in Sentinel 001. "
     "NOTE: A60 is defined in the FTP Code / SOLAS, not in this chapter. Cross-reference required.",
     "Chrystal Revis (Fire/Safety)", "HIGH"],

    ["F-03", "FTP Code",
     "International Code for Application of Fire Test Procedures (2010 edition). "
     "Governing test standard for all habitat interior materials.",
     "Sec.8 [3.1] et al.",
     "M-02 (Alcantara) and M-04 (Teak) rejected. Substitutes M-02S and M-04S must pass FTP Code testing. "
     "All interior material qualification references this code.",
     "Chrystal Revis; Materials programme (BT3345_9001_005)", "HIGH"],

    ["F-04", "non-combustible",
     "Material classification required for thermal insulation, structural elements, "
     "furniture frames (case and free-standing), and certain wall/ceiling surfaces.",
     "Sec.8 [3.4], [3.8], [3.9]",
     "All Sentinel 001 interior materials must meet non-combustible or flame-retardant per FTP Code. "
     "Applies to M-02S and M-04S substitute qualification.",
     "Materials programme (BT3345_9001_005)", "HIGH"],

    ["F-05", "flame-retardant",
     "Material classification tested per FTP Code. Required for room dividers, ceilings, "
     "furniture, textiles, floor coverings, draperies, upholstery, and bedding.",
     "Sec.8 [3.1], [3.5]-[3.14]",
     "Minimum acceptable classification for Sentinel 001 interior materials where "
     "non-combustible is not achievable.",
     "Materials programme (BT3345_9001_005)", "HIGH"],

    ["F-06", "SAT electrical de-energization",
     "SAT-specific: electrical equipment that cannot be assessed per general fire requirements "
     "shall automatically de-energize before atmosphere can support combustion at worst-case ppO2. "
     "Shall happen automatically before chamber atmosphere supports combustion.",
     "Sec.8 [4.1]",
     "Automatic de-energization circuit required in Sentinel 001 -- triggered at elevated ppO2 threshold. "
     "Medical equipment overrides permitted only with risk assessment and auto-reset.",
     "Oguzhan Kilic (Electrical) -- design requirement", "HIGH"],
]

PART_G_HEADERS = ["ID", "Term", "Definition / Requirement", "Source Clause",
                  "Sentinel 001 Application", "Owner", "Confidence"]
PART_G_WIDTHS  = [0.40, 1.00, 2.40, 0.80, 2.90, 1.30, 0.70]

PART_G_DATA = [
    ["G-01", "survival time",
     "Minimum time the habitat must sustain life support without external support, consistent with "
     "emergency rescue concept and not less than minimum per Sec.2. Drives emergency power, gas, "
     "and CO2 scrubber endurance requirements.",
     "Sec.9 [3]; Sec.2",
     "Sentinel 001: 120-hr battery endurance cited as emergency basis in programme context. "
     "Survival time must be formally defined in rescue concept. See Flag 5 and Flag 7.",
     "All leads. Oguzhan Kilic (Electrical) for battery basis.", "HIGH"],

    ["G-02", "emergency rescue concept",
     "Document defining rescue scenarios, responsibilities, equipment, systems, and escape routes. "
     "Required documentation type G170 (For Information). Includes contingency plans.",
     "Sec.9 [2]; Doc Table 12 (G170)",
     "Required AiP document -- not confirmed as submitted in current programme record. "
     "Rescue concept drives survival time, hyperbaric evacuation vehicle design, and place of safety.",
     "Joe Reeves (PM) -- status confirmation required", "HIGH"],

    ["G-03", "hyperbaric evacuation vehicle",
     "Vehicle for simultaneous rescue of all habitat occupants at nominal diving depth. "
     "Requirements: autonomous LSS; rapid-couple connections for gas/water/electrical; "
     "viewports; internal comms; independent power; mating system; crane attachment points.",
     "Sec.9 [7.1]-[7.10]",
     "REQUIRED for LIVHAB(SAT) -- must accommodate all 9 Sentinel 001 occupants at NDD 13.7 msw. "
     "See Flag 4 for programme status and Vanguard/Tennessee Reef asset context.",
     "Joe Reeves (PM); rescue system lead TBD", "HIGH"],

    ["G-04", "transfer vehicle",
     "Vehicle providing breathing gas supply for all occupants in emergency. "
     "Must carry minimum 24 hours gas supply plus supplemental O2 bottle.",
     "Sec.9 [3.3], [3.4]",
     "Rescue system design TBD -- to be defined in emergency rescue concept.",
     "Jim Williamson / rescue system design", "MEDIUM"],

    ["G-05", "marker buoy",
     "Optional emergency signalling device releasable from inside habitat. "
     "Design pressure 1.1 x CDP. May include automatic emergency call transmitter.",
     "Sec.9 [8.1]-[8.3]",
     "TBD per rescue concept. May serve as telephone connection with rescue forces.",
     "Rescue system option -- per CONOPS", "MEDIUM"],

    ["G-06", "emergency mating flange",
     "Mating flange for docking of rescue vehicle. Design parameters and calculations "
     "shall be agreed with the Society.",
     "Sec.9 [9.1]",
     "Required if hyperbaric evacuation vehicle is specified. Design parameters need "
     "prior DNV agreement -- AiP basis item.",
     "Joseph Welker (Structural) / DNV Hovik", "HIGH"],
]

PART_H_HEADERS = ["ID", "Code", "Full Title", "Type", "Sentinel 001 Status", "Note"]
PART_H_WIDTHS  = [0.40, 0.70, 1.80, 0.70, 3.50, 2.40]

PART_H_DATA = [
    ["H-01", "AP", "For Approval -- submission to DNV for formal review and sign-off.",
     "Submission Category",
     "Required for: stability analysis, detailed hull drawings, design analysis, "
     "single line electrical diagrams, test procedures.",
     "Non-negotiable DNV gateway. AP submissions govern what DNV formally approves."],

    ["H-02", "FI", "For Information -- submitted to DNV for awareness, not formal approval.",
     "Submission Category",
     "Required for: CONOPS, FMEA, GA plan, arrangement plans, operation manual, "
     "emergency rescue concept, human factors assessment.",
     "Must still be submitted on schedule. Withholding FI documents is a programme risk."],

    ["H-03", "CONOPS (Z100)", "Concept of operations statement -- defines tasks, environmental "
     "conditions, and intended modes of operation.",
     "FI",
     "Required per Table 6. Forms basis for all Sentinel 001 design decisions. "
     "Governs qualifier applicability, rescue concept scope, and support asset classification.",
     "AiP prerequisite. Must be issued before dependent design documents are submitted."],

    ["H-04", "FMEA (Z071)", "Failure mode and effect analysis -- minimum two independent analyses; "
     "detectability/probability/criticality RPN basis; risk-lowering measures.",
     "FI",
     "OI-07 OPEN: structural tab 3/10 -- not AiP-ready. Two analyses recommended per Sec.1 [9.1] Table 6. "
     "See Flag 6 for verbal form determination.",
     "Verbal form is 'should' (RECOMMENDED, not 'shall'). However, inadequate FMEA will be "
     "flagged by DNV. Quality issue regardless of mandatory status."],
]

# ── FLAG DATA ─────────────────────────────────────────────────────────────────

FLAGS = [
    {
        "num": "1",
        "severity": "HIGH",
        "title": "ppO2 REGIME BOUNDARY -- ATMOS vs. SAT LIMITS",
        "body": (
            "Sec.5 [2.1] sets ppO2 at 0.19-0.23 bar for ATMOS habitats. Sec.5 [3] overrides this "
            "for SAT mode at 0.19-1.6 bar maximum. Sentinel 001 nominally operates at ppO2 = 0.50 bar "
            "breathing air at 13.7 msw -- this is only within limits under the SAT clause. If any DNV "
            "reviewer applies the ATMOS clause in error, the habitat appears non-compliant. All programme "
            "submissions must explicitly cite Sec.5 [3] as the controlling ppO2 requirement. "
            "Flag for inclusion in TQ-AIPBASIS-001 follow-up."
        ),
        "bg": LIGHT_RED,
    },
    {
        "num": "2",
        "severity": "HIGH",
        "title": "PVHO TERM UNDEFINED IN PT.5 CH.10",
        "body": (
            "The term 'PVHO' (pressure vessel for human occupancy) appears in Sec.2 [3.3] of this chapter "
            "but is not defined in the Sec.1 definitions table. The defining document is Pt.1 Ch.4 and/or "
            "ASME PVHO-1. Any Sentinel 001 programme document that cites PVHO must reference the correct "
            "defining document. Do not rely on Pt.5 Ch.10 as the source for this definition."
        ),
        "bg": LIGHT_RED,
    },
    {
        "num": "3",
        "severity": "HIGH",
        "title": "INNER/OUTER AREA DEFINITIONS NOT IN PT.5 CH.10",
        "body": (
            "'Inner area' and 'outer area' are used in Sec.2 [2.2-2.3] of this chapter but are defined "
            "in DNV-OS-E402 Ch.1 Sec.1 Table 4 -- not in Pt.5 Ch.10 itself. SDD-SENTINEL-001 and all "
            "Sentinel 001 deliverables that reference inner or outer area must cite DNV-OS-E402 as the "
            "defining authority."
        ),
        "bg": LIGHT_RED,
    },
    {
        "num": "4",
        "severity": "HIGH",
        "title": "HYPERBARIC EVACUATION VEHICLE -- NOT CONFIRMED IN PROGRAMME RECORD",
        "body": (
            "Sec.9 [7.1] is an explicit SAT requirement for LIVHAB(SAT) systems. A hyperbaric evacuation "
            "vehicle capable of holding all 9 occupants simultaneously at the nominal diving depth of "
            "13.7 msw is mandatory. This item does not appear as an identified deliverable in the current "
            "programme document suite. "
            "NOTE -- Vanguard infrastructure and existing Tennessee Reef rescue assets may partially or "
            "fully satisfy the hyperbaric evacuation vehicle requirement at Sec.9 [7.1]. Resolution "
            "requires CONOPS confirmation that existing rescue assets can accommodate all 9 Sentinel 001 "
            "occupants at 13.7 msw NDD before this is declared a programme gap vs. a documentation gap. "
            "Do not declare AiP-blocking until CONOPS review is complete."
        ),
        "bg": LIGHT_RED,
    },
    {
        "num": "5",
        "severity": "HIGH",
        "title": "MINIMUM SURVIVAL TIME VALUE NOT STATED IN PT.5 CH.10",
        "body": (
            "This document requires survival time to be 'consistent with the emergency rescue concept, "
            "but not less than the minimum.' No absolute minimum value is stated in the text of "
            "Pt.5 Ch.10. The minimum is likely defined in a referenced standard (DNV-OS-E402 or USN "
            "Diving Manual Rev 7). The 120-hr battery endurance cited in the programme context may not "
            "equal the required survival time -- these are separate parameters. Clarification required "
            "from DNV or supplementary clause review before endurance figures are cited in submittals. "
            "See also Flag 7."
        ),
        "bg": LIGHT_RED,
    },
    {
        "num": "6",
        "severity": "RECOMMENDED (not mandatory)",
        "title": "FMEA DUAL ANALYSIS REQUIREMENT -- VERBAL FORM IS 'SHOULD'",
        "body": (
            "NOTE 1 APPLIED: Exact DNV verbal form verified from Sec.1 [9.1] Table 6, Z071 entry. "
            "Text reads: 'At least two independent analyses should be carried out, each with a "
            "determination of the risk based on detectability, probability, and criticality, as well "
            "as the derivation of risk lowering measures. The following should be considered: failure "
            "at component level (single failure); specification of critical scenarios (fire, water "
            "intrusion, etc.).' "
            "Per DNV verbal form hierarchy: 'should' = RECOMMENDED, not mandatory (not 'shall'). "
            "Severity is therefore RECOMMENDED, not BLOCKING. "
            "However: OI-07 structural FMEA tab scored 3/10 remains an active open item. DNV will "
            "assess FMEA quality regardless of mandatory status. An inadequate FMEA will generate "
            "observations or conditions at survey regardless of this verbal form distinction. "
            "Owner: Ken Ngo."
        ),
        "bg": LIGHT_YELLOW,
    },
    {
        "num": "7",
        "severity": "HIGH",
        "title": "SURVIVAL TIME / BATTERY ENDURANCE INTERNAL PROGRAMME CONFLICT",
        "body": (
            "NOTE 2 APPLIED. Internal programme conflict identified in endurance figures: "
            "SDD-SENTINEL-001 Sec 5.3 states 80-hour emergency endurance at 250W critical load. "
            "Programme context states 120-hour endurance. Battery is 60 kWh total (2x30 kWh LiFePO4). "
            "Check: 60 kWh / 250W = 240 hr theoretical maximum. "
            "The 80-hour figure implies the actual critical load is approximately 750W, not 250W -- "
            "OR the battery capacity basis used in SDD-SENTINEL-001 differs from the 60 kWh programme "
            "parameter. These figures cannot both be correct on the same battery and load assumptions. "
            "Resolution required before any DNV submittal cites endurance figures. The survival time "
            "documentation submitted to DNV must be internally consistent across all documents. "
            "Owner: Oguzhan Kilic (Electrical). Confidence: HIGH."
        ),
        "bg": LIGHT_RED,
    },
]


# ── MAIN BUILD ────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Remove default Normal style spacing
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(4)

    setup_page(doc)
    setup_hf(doc)

    # ── COVER / TITLE BAR ────────────────────────────────────────────────
    add_title_bar(doc)
    add_header_block(doc)

    doc.add_paragraph()
    add_body_para(doc,
        "This document defines and maps all terminology drawn from DNV-RU-UWT Pt.5 Ch.10 "
        "(December 2024) to the Sentinel 001 programme context. Definitions are sourced "
        "exclusively from the referenced DNV document unless explicitly noted otherwise. "
        "Programme compliance flags and confidence levels are assigned by Bastion Technologies, "
        "Inc. engineering staff. This document forms part of the Sentinel 001 AiP documentation "
        "suite and is designated programme document number BT3345_9001_008.",
        before=80, after=120)

    # ── PART A ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART A -- CLASS NOTATIONS AND QUALIFIERS")
    add_source_note(doc, "Source: Sec.1 [6], Table 1 | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_A_HEADERS, PART_A_DATA, PART_A_WIDTHS)

    # ── PART B ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART B -- DEFINED TERMS")
    add_source_note(doc, "Source: Sec.1 [8.1], Table 4 | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_B_HEADERS, PART_B_DATA, PART_B_WIDTHS)

    # ── PART C ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART C -- ABBREVIATIONS")
    add_source_note(doc, "Source: Sec.1 [8.2], Table 5 | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_C_HEADERS, PART_C_DATA, PART_C_WIDTHS)

    # ── PART D ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART D -- STRUCTURAL AND PRESSURE DESIGN TERMS")
    add_source_note(doc, "Source: Sec.3 [1.7] Tables 1 & 2; Sec.2 [2.2-2.3] | DNV-OS-E402 Ch.1 Sec.1 Table 4 (for inner/outer area)")
    make_table(doc, PART_D_HEADERS, PART_D_DATA, PART_D_WIDTHS)

    # ── PART E ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART E -- ATMOSPHERIC AND LIFE SUPPORT LIMITS")
    add_source_note(doc, "Source: Sec.5 [1.5], [2.1]-[2.3], [3] | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_E_HEADERS, PART_E_DATA, PART_E_WIDTHS)

    # ── PART F ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART F -- FIRE PREVENTION AND PROTECTION TERMS")
    add_source_note(doc, "Source: Sec.8 [2.1], [3.1]-[3.14], [4.1] | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_F_HEADERS, PART_F_DATA, PART_F_WIDTHS)

    # ── PART G ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART G -- RESCUE AND EMERGENCY SYSTEM TERMS")
    add_source_note(doc, "Source: Sec.9 [2]-[9] | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_G_HEADERS, PART_G_DATA, PART_G_WIDTHS)

    # ── PART H ────────────────────────────────────────────────────────────
    add_section_heading(doc, "PART H -- DOCUMENTATION FRAMEWORK TERMS")
    add_source_note(doc, "Source: Sec.1 [9.1], Tables 6-12 | DNV-RU-UWT Pt.5 Ch.10 (December 2024)")
    make_table(doc, PART_H_HEADERS, PART_H_DATA, PART_H_WIDTHS)

    # ── CRITICAL FLAGS ────────────────────────────────────────────────────
    add_section_heading(doc, "CRITICAL PROGRAMME FLAGS", before=200)
    add_source_note(doc,
        "Flags below identify issues, gaps, and conflicts identified during matrix construction. "
        "HIGH (red): requires resolution before DNV submittal. "
        "RECOMMENDED (yellow): advisory finding -- action recommended but not mandatory per DNV verbal form. "
        "All flags remain open until closed by programme leads.")

    bg_map = {
        LIGHT_RED:    LIGHT_RED,
        LIGHT_YELLOW: LIGHT_YELLOW,
        LIGHT_GREEN:  LIGHT_GREEN,
    }

    for flag in FLAGS:
        add_flag_box(doc,
                     flag_num=flag["num"],
                     severity=flag["severity"],
                     title=flag["title"],
                     body=flag["body"],
                     bg_hex=flag["bg"])

    # ── DOCUMENT CONTROL ─────────────────────────────────────────────────
    doc.add_paragraph()
    add_doc_control(doc)

    # ── SAVE ──────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
